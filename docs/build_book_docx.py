#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AppleWatchDiet_100methods.md → Kindle対応のおしゃれな .docx を生成する。

特徴:
- 見出しスタイル(Heading 1/2/3)を使用 → KDPがナビゲーション目次を自動生成
- 「目次」のリンクをブックマークへの内部ハイパーリンクに変換 → 本文内でもタップ移動可能
- 配色: 黄緑 / 優しいオレンジ / 薄いブルー を基調にしたおしゃれなデザイン
- 画像は縮小して埋め込み（ファイルサイズ削減）
"""

import os
import re
import tempfile

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

SRC = os.path.join(os.path.dirname(__file__), "AppleWatchDiet_100methods.md")
OUT = os.path.join(os.path.dirname(__file__), "AppleWatchDiet_100methods_PB.docx")
BASE_DIR = os.path.dirname(__file__)

# ── 配色パレット ─────────────────────────────────────────────
GREEN       = "7CB342"   # 黄緑（メインカラー: タイトル・大見出し）
GREEN_DK    = "689F38"
GREEN_LT    = "E8F3DA"   # 黄緑の薄い背景
ORANGE      = "EF9A4D"   # 優しいオレンジ（中見出し）
ORANGE_DK   = "E0822F"
ORANGE_LT   = "FCEBD9"   # オレンジの薄い背景
BLUE        = "4FA3D1"   # 薄いブルー（小見出し・リンク）
BLUE_DK     = "3B8AB8"
BLUE_LT     = "E2F2FB"   # ブルーの薄い背景
INK         = "37474F"   # 本文（やわらかいダークスレート）
GRAY_LT     = "F2F3F4"   # コードブロック背景
JP_FONT     = "Hiragino Sans"

_bm_id = [1000]


# ── XMLヘルパー ──────────────────────────────────────────────
def set_style_font(style, *, name=JP_FONT, size=None, color=None, bold=None):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def para_border(p, *, edges=("left",), color=GREEN, sz="18", space="10"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), space)
        el.set(qn("w:color"), color)
        pBdr.append(el)
    pPr.append(pBdr)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_bookmark_to_paragraph(p, name):
    _bm_id[0] += 1
    bid = str(_bm_id[0])
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    p._p.insert(0, start)
    p._p.append(end)


def add_internal_hyperlink(p, anchor, text, *, color=BLUE_DK, bold=False):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), JP_FONT)
    rPr.append(rfonts)
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    if bold:
        rPr.append(OxmlElement("w:b"))
    run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)


def anchor_to_bm(anchor):
    return "bm_" + anchor.strip().lstrip("#").replace("-", "_")


# ── インライン記法 ───────────────────────────────────────────
INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|\[[^\]]+\]\([^)]+\))")


def add_inline(p, text, *, base_color=INK):
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
            r.font.color.rgb = RGBColor.from_string(base_color)
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
            r.font.color.rgb = RGBColor.from_string(base_color)
        elif tok.startswith("[") and "](" in tok:
            label = tok[1:tok.index("](")]
            target = tok[tok.index("](") + 2:-1]
            if target.startswith("#"):
                add_internal_hyperlink(p, anchor_to_bm(target), label)
            else:
                r = p.add_run(label)
                r.font.color.rgb = RGBColor.from_string(BLUE_DK)
        else:
            r = p.add_run(tok)
            r.font.color.rgb = RGBColor.from_string(base_color)


# ── 画像縮小 ─────────────────────────────────────────────────
_tmpdir = tempfile.mkdtemp(prefix="book_imgs_")
_img_cache = {}


def prep_image(rel_path):
    if rel_path in _img_cache:
        return _img_cache[rel_path]
    src = os.path.join(BASE_DIR, rel_path)
    if not os.path.isfile(src):
        _img_cache[rel_path] = None
        return None
    try:
        im = Image.open(src)
        im = im.convert("RGB")
        long_side = 1100
        w, h = im.size
        scale = min(1.0, long_side / max(w, h))
        if scale < 1.0:
            im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        out = os.path.join(_tmpdir, f"img_{len(_img_cache)}.jpg")
        im.save(out, "JPEG", quality=82, optimize=True)
        result = (out, im.size)
        _img_cache[rel_path] = result
        return result
    except Exception as e:
        print("  image skip:", rel_path, e)
        _img_cache[rel_path] = None
        return None


def add_image(doc, rel_path):
    prepped = prep_image(rel_path)
    if not prepped:
        return
    path, (w, h) = prepped
    max_w, max_h = 3.5, 4.7  # inches
    if (w / max_w) >= (h / max_h):
        kwargs = {"width": Inches(max_w)}
    else:
        kwargs = {"height": Inches(max_h)}
    doc.add_picture(path, **kwargs)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(4)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(10)


# ── 各ブロックビルダー ───────────────────────────────────────
def add_heading(doc, level, text, pending_anchor, in_toc):
    # 目次内の "### カテゴリ" は見出しスタイルにせず、色付き段落にする（ナビ重複防止）
    if in_toc and level == 3:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor.from_string(ORANGE_DK)
        return None

    style_map = {1: "Title", 2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}
    p = doc.add_paragraph(style=style_map[level])
    add_inline(p, text, base_color={
        1: GREEN_DK, 2: GREEN_DK, 3: ORANGE_DK, 4: BLUE_DK
    }[level])
    # 大見出し(##)は下線ボーダーで装飾
    if level == 2:
        para_border(p, edges=("bottom",), color=GREEN, sz="8", space="6")
    if pending_anchor:
        add_bookmark_to_paragraph(p, anchor_to_bm(pending_anchor))
    return None


def add_blockquote(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    shade_paragraph(p, BLUE_LT)
    para_border(p, edges=("left",), color=BLUE, sz="24", space="10")
    for i, ln in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        add_inline(p, ln, base_color=INK)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_inline(p, text, base_color=INK)
    # 箇条書きマーカーを黄緑に
    pPr = p._p.get_or_add_pPr()
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), GREEN_DK); rpr.append(c)
    pPr.append(rpr)


def add_paragraph_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    add_inline(p, text, base_color=INK)


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    para_border(p, edges=("bottom",), color=GREEN, sz="6", space="1")


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Pt(8)
    shade_paragraph(p, GRAY_LT)
    para_border(p, edges=("left",), color=GREEN, sz="18", space="10")
    for i, ln in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(ln)
        r.font.name = "Menlo"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string("455A64")
        rpr = r._element.get_or_add_rPr()
        rf = rpr.get_or_add_rFonts()
        rf.set(qn("w:ascii"), "Menlo")
        rf.set(qn("w:hAnsi"), "Menlo")


def add_table(doc, rows):
    # rows: list of list[str]; rows[0]=header, rows[1]=separator(除外済み)
    header = rows[0]
    body = rows[1:]
    ncol = len(header)
    table = doc.add_table(rows=1 + len(body), cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # ヘッダー
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        shade_cell(cell, GREEN)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(cell_text.strip())
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.size = Pt(10)
    # 本文
    for i, brow in enumerate(body):
        for j in range(ncol):
            cell = table.rows[i + 1].cells[j]
            shade_cell(cell, GREEN_LT if i % 2 == 0 else "FFFFFF")
            cell.paragraphs[0].text = ""
            txt = brow[j].strip() if j < len(brow) else ""
            add_inline(cell.paragraphs[0], txt, base_color=INK)
            cell.paragraphs[0].runs and setattr(cell.paragraphs[0].runs[0].font, "size", Pt(9.5))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def split_table_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c for c in s.split("|")]


# ── メイン処理 ───────────────────────────────────────────────
def build():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()

    # ベーススタイル
    set_style_font(doc.styles["Normal"], name=JP_FONT, size=10.5, color=INK)
    set_style_font(doc.styles["Title"], name=JP_FONT, size=26, color=GREEN_DK, bold=True)
    set_style_font(doc.styles["Heading 1"], name=JP_FONT, size=19, color=GREEN_DK, bold=True)
    set_style_font(doc.styles["Heading 2"], name=JP_FONT, size=14.5, color=ORANGE_DK, bold=True)
    set_style_font(doc.styles["Heading 3"], name=JP_FONT, size=12, color=BLUE_DK, bold=True)
    for sname in ("List Bullet",):
        try:
            set_style_font(doc.styles[sname], name=JP_FONT, size=10.5, color=INK)
        except KeyError:
            pass

    pending_anchor = None
    in_toc = False
    in_code = False
    code_lines = []

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # コードフェンス
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # アンカー
        m = re.match(r'<a id="([^"]+)"></a>', stripped)
        if m:
            pending_anchor = m.group(1)
            i += 1
            continue

        # 空行
        if stripped == "":
            i += 1
            continue

        # 水平線
        if stripped == "---":
            add_divider(doc)
            i += 1
            continue

        # 見出し
        hm = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            level = min(level, 4)
            if level == 2 and text == "目次":
                in_toc = True
            elif level <= 2 and text != "目次":
                in_toc = False
            add_heading(doc, level, text, pending_anchor, in_toc)
            pending_anchor = None
            i += 1
            continue

        # 画像
        im = re.match(r"^!\[[^\]]*\]\(([^)]+)\)\s*$", stripped)
        if im:
            add_image(doc, im.group(1))
            i += 1
            continue

        # テーブル
        if stripped.startswith("|"):
            tbl = []
            while i < n and lines[i].strip().startswith("|"):
                tbl.append(lines[i])
                i += 1
            parsed = [split_table_row(r) for r in tbl]
            # セパレータ行(---)を除去
            parsed = [r for r in parsed if not all(set(c.strip()) <= set("-: ") and c.strip() for c in r)]
            if parsed:
                add_table(doc, parsed)
            continue

        # 引用
        if stripped.startswith(">"):
            quote = []
            while i < n and lines[i].strip().startswith(">"):
                q = lines[i].strip()[1:].strip()
                quote.append(q.rstrip())
                i += 1
            add_blockquote(doc, quote)
            continue

        # 箇条書き
        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:].strip())
            i += 1
            continue

        # 通常段落
        add_paragraph_text(doc, stripped)
        i += 1

    doc.save(OUT)
    size_mb = os.path.getsize(OUT) / 1048576
    print(f"✅ Saved: {OUT}  ({size_mb:.1f} MB)")


if __name__ == "__main__":
    build()
