#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
cursor-claude-code-ios-app-book-plus.md → Kindle Paperback 対応 .docx 生成スクリプト

特徴:
- 用紙サイズ: 210 × 257 mm (Kindle Paperback 8.27 × 10.11 インチ)
- 表紙ページ（ページ番号なし）
- Heading 1/2/3 スタイル → KDP ナビゲーション目次を自動生成
- フッターにページ番号（中央揃え）
- <small> タグ → 8pt グレー文字で免責事項を描画
- Plus テーマカラー（パープル × オレンジ × グリーン）
- コードブロック / テーブル / 引用 / 画像 対応
- `[プロンプト例]:` ラベル付きコードブロックを枠付きで強調

pip install python-docx Pillow
"""

import os
import re
import tempfile

from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

SRC = os.path.join(os.path.dirname(__file__), "cursor-claude-code-ios-app-book-plus.md")
OUT = os.path.join(os.path.dirname(__file__), "cursor-claude-code-ios-app-book-plus.docx")
BASE_DIR = os.path.dirname(__file__)

# ── Kindle Paperback ページサイズ ──────────────────────────────
PAGE_W        = Mm(210)
PAGE_H        = Mm(257)
MARGIN_TOP    = Mm(19)
MARGIN_BOTTOM = Mm(19)
MARGIN_INNER  = Mm(19)
MARGIN_OUTER  = Mm(13)

# ── 配色パレット（Fitingo Plus テーマ）────────────────────────
GREEN      = "58CC02"   # Fitingo グリーン（H1・章タイトル）
GREEN_DK   = "3E8F00"
GREEN_LT   = "EDFAD4"
ORANGE     = "FF8C00"   # Fitingo オレンジ（Plus カラー・H2）
ORANGE_DK  = "CC6F00"
ORANGE_LT  = "FFF4E0"
PURPLE     = "9747E8"   # Plus パープル（H3）
PURPLE_DK  = "7228C0"
PURPLE_LT  = "F3EBFE"
BLUE_DK    = "3B8AB8"
INK        = "37474F"   # 本文
GRAY_LT    = "F2F3F4"   # コードブロック背景
CODE_HINT  = "FFF8EC"   # プロンプト例背景（薄いオレンジ）
JP_FONT    = "Hiragino Sans"

_bm_id = [2000]

# TOC に載せない見出しテキスト（Heading 1 スタイルを使わずに手動フォーマット）
HEADING_NO_TOC = {"免責事項・著作権表示"}


# ── XML ヘルパー ──────────────────────────────────────────────
def set_style_font(style, *, name=JP_FONT, size=None, color=None, bold=None):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def ensure_style_outline_level(style, level):
    """スタイルの pPr に w:outlineLvl を設定（KDP が heading として認識するために必須）"""
    pPr = style.element.get_or_add_pPr()
    for el in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(el)
    ol = OxmlElement("w:outlineLvl")
    ol.set(qn("w:val"), str(level))
    pPr.append(ol)


def set_paragraph_outline_level(p, level):
    """段落の pPr に w:outlineLvl を明示設定（Kindle NCX ナビゲーション用）"""
    pPr = p._p.get_or_add_pPr()
    for el in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(el)
    ol = OxmlElement("w:outlineLvl")
    ol.set(qn("w:val"), str(level))
    pPr.append(ol)


def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def para_border(p, *, edges=("left",), color=GREEN, sz="18", space="10"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), space)
        el.set(qn("w:color"), color)
        pBdr.append(el)
    pPr.append(pBdr)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_bookmark(p, name):
    _bm_id[0] += 1
    bid = str(_bm_id[0])
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    p._p.insert(0, start)
    p._p.append(end)


def add_internal_hyperlink(p, anchor, text, *, color=BLUE_DK, bold=False):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), JP_FONT)
    rPr.append(rfonts)
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    if bold:
        rPr.append(OxmlElement("w:b"))
    run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)


def anchor_to_bm(anchor):
    return "bm_" + anchor.strip().lstrip("#").replace("-", "_")


# ── インライン記法 ────────────────────────────────────────────
INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_inline(p, text, *, base_color=INK):
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
            r.font.color.rgb = RGBColor.from_string(base_color)
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
            r.font.color.rgb = RGBColor.from_string(base_color)
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1])
            r.font.name = "Menlo"
            r.font.color.rgb = RGBColor.from_string(PURPLE_DK)
        elif tok.startswith("[") and "](" in tok:
            label = tok[1:tok.index("](")]
            target = tok[tok.index("](") + 2:-1]
            if target.startswith("#"):
                add_internal_hyperlink(p, anchor_to_bm(target), label)
            else:
                r = p.add_run(label)
                r.font.color.rgb = RGBColor.from_string(BLUE_DK)
        else:
            r = p.add_run(tok)
            r.font.color.rgb = RGBColor.from_string(base_color)


# ── 画像縮小キャッシュ ────────────────────────────────────────
_tmpdir = tempfile.mkdtemp(prefix="plus_book_imgs_")
_img_cache = {}


def prep_image(rel_path):
    if rel_path in _img_cache:
        return _img_cache[rel_path]
    src = os.path.join(BASE_DIR, rel_path)
    if not os.path.isfile(src):
        print(f"  ⚠️  画像が見つかりません: {rel_path}")
        _img_cache[rel_path] = None
        return None
    try:
        im = Image.open(src).convert("RGB")
        long_side = 1100
        w, h = im.size
        scale = min(1.0, long_side / max(w, h))
        if scale < 1.0:
            im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        out = os.path.join(_tmpdir, f"img_{len(_img_cache)}.jpg")
        im.save(out, "JPEG", quality=82, optimize=True)
        result = (out, im.size)
        _img_cache[rel_path] = result
        return result
    except Exception as e:
        print(f"  ⚠️  画像スキップ: {rel_path} ({e})")
        _img_cache[rel_path] = None
        return None


def add_image(doc, rel_path, alt=""):
    prepped = prep_image(rel_path)
    if not prepped:
        return
    path, (w, h) = prepped
    max_w, max_h = 3.8, 3.4
    if (w / max_w) >= (h / max_h):
        kwargs = {"width": Inches(max_w)}
    else:
        kwargs = {"height": Inches(max_h)}
    shape = doc.add_picture(path, **kwargs)
    if alt:
        docPr = shape._inline.find(qn("wp:docPr"))
        if docPr is not None:
            docPr.set("descr", alt)
            docPr.set("title", alt)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(4)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)
    if alt:
        cap = doc.add_paragraph(alt)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        for r in cap.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string("888888")
            r.italic = True


# ── TOC プリスキャン ──────────────────────────────────────────
def prescan_toc(lines):
    """
    Kindle NCX 対応: level 2（章）・3（節）・4（項）を収集。
    TOC 表示フィールドには 2/3 のみ使用。
    level 4 は Kindle ナビゲーション専用（outline level 設定のみ）。
    """
    entries = []
    in_toc = False
    in_code = False
    in_small = False
    counter = 0
    for raw in lines:
        s = raw.strip()
        if s.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        if s.startswith("<small"):
            in_small = True; continue
        if s.startswith("</small"):
            in_small = False; continue
        if in_small:
            continue
        hm = re.match(r"^(#{1,6})\s+(.*)$", s)
        if not hm:
            continue
        level = min(len(hm.group(1)), 4)
        text = hm.group(2).strip()
        if level == 2 and text == "目次":
            in_toc = True
            continue
        if level <= 2 and text != "目次":
            in_toc = False
        if in_toc:
            continue
        clean = _strip_inline(text)
        # 番号なし level 3 見出し（付録サブ節など）はTOCに載せない
        # X-Y. 形式 または X-Y. 末尾b形式 または X-A. 形式（1-A. 等）を有効とする
        is_numbered = bool(re.match(r'^\d+-[\dA-Z]', clean))
        if level == 3 and not is_numbered:
            continue
        if level in (2, 3, 4) and text != "目次" and clean not in HEADING_NO_TOC:
            counter += 1
            entries.append((level, clean, f"_Toc{91000000 + counter}"))
    return entries


def _strip_inline(text):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("**", "").replace("*", "").replace("`", "").strip()


def _toc_field_run(parent_par, char_type=None, instr=None):
    r = parent_par.add_run()
    if char_type:
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), char_type)
        r._r.append(fc)
    if instr is not None:
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = instr
        r._r.append(it)
    return r


def _toc_pageref_run(tocname):
    """TOC エントリのページ番号参照フィールド（{ PAGEREF bookmark \\h }）を生成"""
    # begin
    r_begin = OxmlElement("w:r")
    rPr_b = OxmlElement("w:rPr")
    ncs = OxmlElement("w:noProof"); rPr_b.append(ncs)
    r_begin.append(rPr_b)
    fc_begin = OxmlElement("w:fldChar"); fc_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(fc_begin)
    # instr
    r_instr = OxmlElement("w:r")
    rPr_i = OxmlElement("w:rPr")
    ncs2 = OxmlElement("w:noProof"); rPr_i.append(ncs2)
    r_instr.append(rPr_i)
    instr_t = OxmlElement("w:instrText")
    instr_t.set(qn("xml:space"), "preserve")
    instr_t.text = f" PAGEREF {tocname} \\h "
    r_instr.append(instr_t)
    # separate
    r_sep = OxmlElement("w:r")
    rPr_s = OxmlElement("w:rPr")
    ncs3 = OxmlElement("w:noProof"); rPr_s.append(ncs3)
    r_sep.append(rPr_s)
    fc_sep = OxmlElement("w:fldChar"); fc_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fc_sep)
    # placeholder value
    r_val = OxmlElement("w:r")
    rPr_v = OxmlElement("w:rPr")
    sz_v = OxmlElement("w:sz"); sz_v.set(qn("w:val"), "20"); rPr_v.append(sz_v)
    r_val.append(rPr_v)
    t_val = OxmlElement("w:t"); t_val.text = "1"; r_val.append(t_val)
    # end
    r_end = OxmlElement("w:r")
    rPr_e = OxmlElement("w:rPr")
    ncs4 = OxmlElement("w:noProof"); rPr_e.append(ncs4)
    r_end.append(rPr_e)
    fc_end = OxmlElement("w:fldChar"); fc_end.set(qn("w:fldCharType"), "end")
    r_end.append(fc_end)
    return [r_begin, r_instr, r_sep, r_val, r_end]


def _add_toc_tab(p, level):
    """TOC エントリの pPr にドットリーダー付き右揃えタブストップを設定する。
    タブ文字の挿入は emit_toc_field 内で hyperlink の直後に行う。"""
    pPr = p._p.get_or_add_pPr()
    tabs_el = pPr.find(qn("w:tabs"))
    if tabs_el is None:
        tabs_el = OxmlElement("w:tabs")
        pPr.append(tabs_el)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    # 右端タブ位置（twips）: 210mm - 19mm(内側) - 13mm(外側) = 178mm ≈ 10080twips
    # level 2（章）: インデントなし、level 3（節）: 18pt インデント
    tab_pos = "9640" if level == 2 else "9360"
    tab.set(qn("w:pos"), tab_pos)
    tabs_el.append(tab)


def emit_toc_field(doc, entries):
    """
    可視 TOC フィールドを生成。
    各エントリにドットリーダー＋PAGEREF ページ番号フィールドを追加。
    \\o "1-2"  → Heading 1/2（章・節）のみ表示
    \\h        → ハイパーリンク
    \\z        → Web レイアウトではタブを非表示
    Kindle NCX は見出しスタイルから別途自動生成。
    """
    if not entries:
        return
    # 可視 TOC には level 2・3（章・節）のみ表示
    visible = [(lv, tx, bm) for (lv, tx, bm) in entries if lv in (2, 3)]
    if not visible:
        return
    last = len(visible) - 1
    for idx, (level, text, tocname) in enumerate(visible):
        ep = doc.add_paragraph()
        ep.paragraph_format.left_indent = Pt(0 if level == 2 else 18)
        ep.paragraph_format.space_after = Pt(3)
        ep.paragraph_format.line_spacing = 1.1
        # ドットリーダータブをパラグラフに設定
        _add_toc_tab(ep, level)

        if idx == 0:
            _toc_field_run(ep, char_type="begin")
            _toc_field_run(ep, instr=' TOC \\o "1-2" \\h \\z ')
            _toc_field_run(ep, char_type="separate")

        # ─ タイトルテキスト（ハイパーリンク）─
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("w:anchor"), tocname)
        hl.set(qn("w:history"), "1")
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts"); rf.set(qn("w:eastAsia"), JP_FONT); rPr.append(rf)
        col = OxmlElement("w:color")
        col.set(qn("w:val"), GREEN_DK if level == 2 else INK)
        rPr.append(col)
        if level == 2:
            rPr.append(OxmlElement("w:b"))
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "26" if level == 2 else "22"); rPr.append(sz)
        run.append(rPr)
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
        run.append(t)
        hl.append(run)
        ep._p.append(hl)

        # ─ タブ文字（ドットリーダーへ） ─
        r_tab = OxmlElement("w:r")
        rPr_tab = OxmlElement("w:rPr")
        sz_tab = OxmlElement("w:sz"); sz_tab.set(qn("w:val"), "22"); rPr_tab.append(sz_tab)
        r_tab.append(rPr_tab)
        t_tab = OxmlElement("w:tab")
        r_tab.append(t_tab)
        ep._p.append(r_tab)

        # ─ ページ番号 PAGEREF フィールド ─
        for r in _toc_pageref_run(tocname):
            ep._p.append(r)

        if idx == last:
            _toc_field_run(ep, char_type="end")


# ── ブロックビルダー ──────────────────────────────────────────
def add_heading(doc, level, text, pending_anchor, page_break=False, toc_bookmark=None):
    """
    見出しを追加する。
    - HEADING_NO_TOC に含まれるテキストは Normal スタイル＋手動フォーマット
      → Word の TOC フィールド（F9 更新時）に拾われない
    - それ以外は Heading 1/2/3 スタイル + outlineLvl（Kindle NCX 対応）
    """
    stripped_text = _strip_inline(text)
    use_no_toc = (level == 2 and stripped_text in HEADING_NO_TOC)

    if use_no_toc:
        # TOC に出ない「章レベル」見出し: Normal スタイルで手動フォーマット
        p = doc.add_paragraph(style="Normal")
        if page_break:
            p.paragraph_format.page_break_before = True
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(8)
        add_inline(p, text, base_color=GREEN_DK)
        for r in p.runs:
            r.font.size = Pt(18)
            r.font.bold = True
            rpr = r._element.get_or_add_rPr()
            rf = rpr.get_or_add_rFonts()
            rf.set(qn("w:eastAsia"), JP_FONT)
        para_border(p, edges=("bottom",), color=GREEN, sz="8", space="6")
        if pending_anchor:
            add_bookmark(p, anchor_to_bm(pending_anchor))
        return None

    # 番号なし level 3 見出し（付録サブ節など）は Heading 3 に格下げ
    # → TOC フィールド \o "1-2" に拾われない（outlineLvl 2 = Heading level 3）
    is_numbered = bool(re.match(r'^\d+-[\dA-Z]', stripped_text))
    if level == 3 and not is_numbered:
        effective_level = 4   # Heading 3 スタイルを使う
    else:
        effective_level = level

    style_map = {1: "Title", 2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}
    # Heading 1/2/3 = outlineLvl 0/1/2（Word/KDP 規約）
    outline_map = {2: 0, 3: 1, 4: 2}
    p = doc.add_paragraph(style=style_map[effective_level])
    if page_break:
        p.paragraph_format.page_break_before = True
    color_map = {1: GREEN_DK, 2: GREEN_DK, 3: ORANGE_DK, 4: PURPLE_DK}
    add_inline(p, text, base_color=color_map[level])
    if level == 2:
        para_border(p, edges=("bottom",), color=GREEN, sz="8", space="6")
    # Kindle NCX: 段落レベルで outlineLvl を明示設定
    if effective_level in outline_map:
        set_paragraph_outline_level(p, outline_map[effective_level])
    if toc_bookmark:
        add_bookmark(p, toc_bookmark)
    if pending_anchor:
        add_bookmark(p, anchor_to_bm(pending_anchor))
    return None


def add_blockquote(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    shade_paragraph(p, ORANGE_LT)
    para_border(p, edges=("left",), color=ORANGE, sz="24", space="10")
    for i, ln in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        add_inline(p, ln, base_color=INK)


def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if indent > 0:
        p.paragraph_format.left_indent = Pt(indent * 18)
    add_inline(p, text, base_color=INK)


def add_numbered(doc, num, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    add_inline(p, text, base_color=INK)


def add_paragraph_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    add_inline(p, text, base_color=INK)


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    para_border(p, edges=("bottom",), color=GREEN, sz="6", space="1")


def add_code_block(doc, lines, is_prompt=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Pt(8)
    shade_paragraph(p, CODE_HINT if is_prompt else GRAY_LT)
    border_color = ORANGE if is_prompt else GREEN
    para_border(p, edges=("left",), color=border_color, sz="18", space="10")
    for i, ln in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(ln)
        r.font.name = "Menlo"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(ORANGE_DK if is_prompt else "455A64")
        rpr = r._element.get_or_add_rPr()
        rf = rpr.get_or_add_rFonts()
        rf.set(qn("w:ascii"), "Menlo")
        rf.set(qn("w:hAnsi"), "Menlo")


def add_table(doc, rows):
    header = rows[0]
    body = rows[1:]
    ncol = len(header)
    table = doc.add_table(rows=1 + len(body), cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        shade_cell(cell, ORANGE)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(cell_text.strip())
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.size = Pt(10)
    for i, brow in enumerate(body):
        for j in range(ncol):
            cell = table.rows[i + 1].cells[j]
            shade_cell(cell, ORANGE_LT if i % 2 == 0 else "FFFFFF")
            cell.paragraphs[0].text = ""
            txt = brow[j].strip() if j < len(brow) else ""
            add_inline(cell.paragraphs[0], txt, base_color=INK)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def split_table_row(line):
    s = line.strip()
    if s.startswith("|"): s = s[1:]
    if s.endswith("|"):   s = s[:-1]
    return [c for c in s.split("|")]


def add_prompt_label(doc, label_text):
    """[プロンプト例]: などのラベル行をオレンジ強調で追加"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label_text)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(ORANGE_DK)


def add_small_text(doc, text):
    """<small> ブロック内のテキストを 8pt グレーで追加"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    add_inline(p, text, base_color="888888")
    for r in p.runs:
        r.font.size = Pt(8)
        rpr = r._element.get_or_add_rPr()
        rf = rpr.get_or_add_rFonts()
        rf.set(qn("w:eastAsia"), JP_FONT)


# ── 用紙サイズ・マージン設定 ──────────────────────────────────
def apply_page_setup(doc):
    for section in doc.sections:
        section.page_width    = PAGE_W
        section.page_height   = PAGE_H
        section.top_margin    = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin   = MARGIN_INNER
        section.right_margin  = MARGIN_OUTER
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)


# ── ページ番号フッター ─────────────────────────────────────────
def _page_num_run(char_type=None, instr=None):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz");   sz.set(qn("w:val"), "20");  rPr.append(sz)
    szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "20"); rPr.append(szCs)
    col = OxmlElement("w:color"); col.set(qn("w:val"), "888888"); rPr.append(col)
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:eastAsia"), JP_FONT); rPr.append(rf)
    r.append(rPr)
    if char_type:
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), char_type)
        r.append(fc)
    if instr is not None:
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = instr
        r.append(it)
    return r


def add_page_number_footer(doc):
    """全セクションのフッター中央にページ番号を追加。表紙（最初のページ）は除外。"""
    for idx, section in enumerate(doc.sections):
        # 表紙ページ（最初のセクション）は先頭ページフッターを空にする
        if idx == 0:
            section.different_first_page_header_footer = True

        footer = section.footer
        footer.is_linked_to_previous = False

        if footer.paragraphs:
            p = footer.paragraphs[0]
            for child in list(p._p):
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag in ("r", "hyperlink", "fldSimple"):
                    p._p.remove(child)
        else:
            p = footer.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(0)

        p._p.append(_page_num_run(char_type="begin"))
        p._p.append(_page_num_run(instr=" PAGE "))
        p._p.append(_page_num_run(char_type="separate"))
        r_val = OxmlElement("w:r")
        rPr2 = OxmlElement("w:rPr")
        sz2 = OxmlElement("w:sz"); sz2.set(qn("w:val"), "20"); rPr2.append(sz2)
        r_val.append(rPr2)
        t_val = OxmlElement("w:t"); t_val.text = "1"; r_val.append(t_val)
        p._p.append(r_val)
        p._p.append(_page_num_run(char_type="end"))


# ── 表紙ページ ────────────────────────────────────────────────
def add_cover_page(doc):
    """表紙ページ: タイトル・サブタイトル・著者名を中央揃えで配置"""
    # 上余白として空行を数行入れる
    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # メインタイトル
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(12)
    tr = tp.add_run("Cursor + Claudeで\n個人開発アプリを\n収益化する方法")
    tr.font.name = JP_FONT
    tr.font.size = Pt(26)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor.from_string(GREEN_DK)
    rpr = tr._element.get_or_add_rPr()
    rf  = rpr.get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), JP_FONT)

    # サブタイトル
    sp2 = doc.add_paragraph()
    sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp2.paragraph_format.space_before = Pt(8)
    sp2.paragraph_format.space_after  = Pt(20)
    sr2 = sp2.add_run("Kindle電子書籍出版 & Plusサブスクリプション実装\n— AI個人開発 完全ガイド —")
    sr2.font.name = JP_FONT
    sr2.font.size = Pt(13)
    sr2.font.bold = False
    sr2.font.color.rgb = RGBColor.from_string(ORANGE_DK)
    rpr2 = sr2._element.get_or_add_rPr()
    rf2  = rpr2.get_or_add_rFonts()
    rf2.set(qn("w:eastAsia"), JP_FONT)

    # 帯（緑の罫線）
    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(4)
    div.paragraph_format.space_after  = Pt(20)
    para_border(div, edges=("bottom",), color=GREEN, sz="12", space="4")

    # 前作との関係を示すキャッチコピー
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after  = Pt(24)
    cr = cp.add_run("『Cursor + ClaudeでiPhoneアプリ・Apple Watchフィットネスアプリを\n週末だけで作る方法』の続編・拡張版")
    cr.font.name = JP_FONT
    cr.font.size = Pt(11)
    cr.font.italic = True
    cr.font.color.rgb = RGBColor.from_string("555555")
    rpr3 = cr._element.get_or_add_rPr()
    rf3  = rpr3.get_or_add_rFonts()
    rf3.set(qn("w:eastAsia"), JP_FONT)

    # 著者名
    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ap.paragraph_format.space_before = Pt(0)
    ap.paragraph_format.space_after  = Pt(6)
    ar = ap.add_run("著者：吉田 顕一（Ken Yoshida）")
    ar.font.name = JP_FONT
    ar.font.size = Pt(12)
    ar.font.bold = True
    ar.font.color.rgb = RGBColor.from_string(INK)
    rpr4 = ar._element.get_or_add_rPr()
    rf4  = rpr4.get_or_add_rFonts()
    rf4.set(qn("w:eastAsia"), JP_FONT)

    # 発行年
    yp = doc.add_paragraph()
    yp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    yp.paragraph_format.space_before = Pt(0)
    yp.paragraph_format.space_after  = Pt(0)
    yr = yp.add_run("2026")
    yr.font.name = JP_FONT
    yr.font.size = Pt(11)
    yr.font.color.rgb = RGBColor.from_string("888888")
    rpr5 = yr._element.get_or_add_rPr()
    rf5  = rpr5.get_or_add_rFonts()
    rf5.set(qn("w:eastAsia"), JP_FONT)

    # 改ページ（表紙の終わり）
    pb = doc.add_paragraph()
    run_pb = pb.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run_pb._r.append(br)


# ── メイン処理 ────────────────────────────────────────────────
def build():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()

    # ベーススタイル設定
    set_style_font(doc.styles["Normal"],    name=JP_FONT, size=10.5, color=INK)
    set_style_font(doc.styles["Title"],     name=JP_FONT, size=24,   color=GREEN_DK,  bold=True)
    set_style_font(doc.styles["Heading 1"], name=JP_FONT, size=18,   color=GREEN_DK,  bold=True)
    set_style_font(doc.styles["Heading 2"], name=JP_FONT, size=14,   color=ORANGE_DK, bold=True)
    set_style_font(doc.styles["Heading 3"], name=JP_FONT, size=12,   color=PURPLE_DK, bold=True)
    for sname in ("List Bullet", "List Number"):
        try:
            set_style_font(doc.styles[sname], name=JP_FONT, size=10.5, color=INK)
        except KeyError:
            pass
    # Kindle NCX: スタイルの pPr に outlineLvl を明示設定
    # Heading 1=0, Heading 2=1, Heading 3=2（Word/KDP の規約）
    ensure_style_outline_level(doc.styles["Heading 1"], 0)
    ensure_style_outline_level(doc.styles["Heading 2"], 1)
    ensure_style_outline_level(doc.styles["Heading 3"], 2)

    # 用紙サイズ・マージン設定（最初のセクション）
    s0 = doc.sections[0]
    s0.page_width         = PAGE_W
    s0.page_height        = PAGE_H
    s0.top_margin         = MARGIN_TOP
    s0.bottom_margin      = MARGIN_BOTTOM
    s0.left_margin        = MARGIN_INNER
    s0.right_margin       = MARGIN_OUTER
    s0.header_distance    = Mm(10)
    s0.footer_distance    = Mm(10)
    # 表紙（先頭ページ）はフッターなし
    s0.different_first_page_header_footer = True

    # 表紙ページを追加
    add_cover_page(doc)

    toc_entries = prescan_toc(lines)
    toc_queue = [e[2] for e in toc_entries]

    pending_anchor = None
    in_toc   = False
    in_code  = False
    in_small = False   # <small> ブロック内フラグ
    code_lines = []
    next_code_is_prompt = False

    # Markdown の H1（# タイトル行）はスキップ（表紙で出力済み）
    skip_first_h1 = True

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # コードフェンス
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines, is_prompt=next_code_is_prompt)
                code_lines = []
                in_code = False
                next_code_is_prompt = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # <small> / </small> ブロック制御
        if stripped.startswith("<small"):
            in_small = True
            i += 1
            continue
        if stripped.startswith("</small"):
            in_small = False
            i += 1
            continue

        # <small> ブロック内: 小文字グレーで描画
        if in_small:
            if stripped and stripped != "---":
                add_small_text(doc, stripped)
            i += 1
            continue

        # <div> タグ系はスキップ
        if stripped.startswith("<div") or stripped.startswith("</div"):
            i += 1
            continue

        # アンカー
        m = re.match(r'<a id="([^"]+)"></a>', stripped)
        if m:
            pending_anchor = m.group(1)
            i += 1
            continue

        # 見出し
        hm = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if hm:
            level = min(len(hm.group(1)), 4)
            text = hm.group(2).strip()

            # H1（# タイトル）は表紙で出力済みのためスキップ
            if level == 1 and skip_first_h1:
                skip_first_h1 = False
                i += 1
                continue

            if level == 2 and text == "目次":
                in_toc = True
                add_heading(doc, level, text, pending_anchor, page_break=True)
                emit_toc_field(doc, toc_entries)
                pending_anchor = None
                i += 1
                continue
            if level <= 2 and text != "目次":
                in_toc = False
            if in_toc:
                pending_anchor = None
                i += 1
                continue
            # すべての ## 見出し（level 2）は改ページ
            # 目次・はじめに・終わりに・付録・各章 すべて新しいページから開始
            is_page_break = (level == 2)
            # level 2/3/4 すべてにブックマーク割り当て（Kindle NCX リンク用）
            toc_bm = None
            if level in (2, 3, 4) and text != "目次":
                toc_bm = toc_queue.pop(0) if toc_queue else None
            add_heading(doc, level, text, pending_anchor,
                        page_break=is_page_break, toc_bookmark=toc_bm)
            pending_anchor = None
            i += 1
            continue

        # 空行
        if stripped == "":
            i += 1
            continue

        # 目次領域の本文はスキップ
        if in_toc:
            i += 1
            continue

        # 水平線
        if stripped == "---":
            add_divider(doc)
            i += 1
            continue

        # [プロンプト例]: ラベル付き行の検出
        label_m = re.match(r"^(\[プロンプト例\][^\:]*):?\s*(.*)$", stripped)
        if label_m:
            label_text = label_m.group(1) + (": " + label_m.group(2) if label_m.group(2) else "")
            add_prompt_label(doc, label_text.strip(": "))
            next_code_is_prompt = True
            i += 1
            continue

        # 画像
        im = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if im:
            add_image(doc, im.group(2), im.group(1))
            i += 1
            continue

        # テーブル
        if stripped.startswith("|"):
            tbl = []
            while i < n and lines[i].strip().startswith("|"):
                tbl.append(lines[i])
                i += 1
            parsed = [split_table_row(r) for r in tbl]
            parsed = [r for r in parsed if not all(set(c.strip()) <= set("-: ") and c.strip() for c in r)]
            if parsed:
                add_table(doc, parsed)
            continue

        # 引用
        if stripped.startswith(">"):
            quote = []
            while i < n and lines[i].strip().startswith(">"):
                q = lines[i].strip()[1:].strip()
                quote.append(q.rstrip())
                i += 1
            add_blockquote(doc, quote)
            continue

        # 箇条書き（インデント対応）
        bm = re.match(r"^(\s*)[-•]\s+(.+)$", line)
        if bm:
            indent = len(bm.group(1)) // 2
            add_bullet(doc, bm.group(2).strip(), indent=indent)
            i += 1
            continue

        # 番号付きリスト
        nm = re.match(r"^\d+\.\s+(.+)$", stripped)
        if nm:
            add_numbered(doc, 0, nm.group(1).strip())
            i += 1
            continue

        # 通常段落
        add_paragraph_text(doc, stripped)
        i += 1

    # ページ番号フッターを全セクションに追加
    add_page_number_footer(doc)

    doc.save(OUT)
    size_mb = os.path.getsize(OUT) / 1048576
    print(f"✅ 保存完了: {OUT}  ({size_mb:.1f} MB)")
    print()
    print("【次のステップ】")
    print("  Word で開いて Ctrl+A → F9 でフィールド更新 → 目次のページ番号が確定します")


if __name__ == "__main__":
    build()
