#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AppleWatchMind_100methods.md → Kindle Paperback 対応 .docx 生成スクリプト

特徴:
- 用紙サイズ: 5.04 × 7.17 インチ (Kindle Paperback トリムサイズ)
- 表紙ページ（ページ番号なし）
- Heading 1/2/3 スタイル → KDP ナビゲーション目次を自動生成
- フッターにページ番号（中央揃え）
- <small> タグ → 8pt グレー文字で免責事項を描画
- マインドフルネステーマカラー（深いティール × 紫 × 青）
- コードブロック / テーブル / 引用 / 画像 対応

pip install python-docx Pillow
"""

import os
import re
import tempfile

from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

SRC = os.path.join(os.path.dirname(__file__), "AppleWatchMind_100methods.md")
OUT = os.path.join(os.path.dirname(__file__), "AppleWatchMind_100methods.docx")
BASE_DIR = os.path.dirname(__file__)

# ── Kindle Paperback ページサイズ（5.04" × 7.17" トリムサイズ）──────
PAGE_W        = Inches(5.04)
PAGE_H        = Inches(7.17)
# KDP推奨マージン（本書は118ページ＝24〜150ページ帯：内側0.375"以上・外側/上下0.25"以上が必要値）
# 内側は最低要件に余裕を持たせ、mirror margins（見開き対応）と併用して奇数/偶数ページどちらでも
# 綴じ側マージンが最低値を下回らないようにする
MARGIN_TOP    = Inches(0.4)
MARGIN_BOTTOM = Inches(0.4)
MARGIN_INNER  = Inches(0.5)
MARGIN_OUTER  = Inches(0.3)

# ── 配色パレット（マインドフルネス・ウェルネステーマ）──────────
# メインカラー：深いティール（安心・瞑想・集中）
TEAL       = "00897B"   # ディープティール（H1・章タイトル・罫線）
TEAL_DK    = "00695C"   # ダークティール（H1/H2 テキスト・テーブルヘッダー）
TEAL_LT    = "E0F2F1"   # ライトティール（ブロック背景・引用）
# アクセントカラー：インディゴブルー（洞察・知性）
INDIGO     = "3949AB"   # インディゴ（H2・アクセント）
INDIGO_DK  = "283593"   # ダークインディゴ（H3・小見出し）
INDIGO_LT  = "E8EAF6"   # ライトインディゴ（コードブロック背景）
# セカンダリ：ラベンダー（穏やかさ・マインドフルネス）
LAVENDER   = "7B1FA2"   # ラベンダー紫（H3）
LAVENDER_DK = "4A148C"  # ダークラベンダー
LAVENDER_LT = "F3E5F5"  # ライトラベンダー
BLUE_DK    = "1565C0"   # アクセントブルー（リンク）
INK        = "263238"   # 本文（ダークグレー）
GRAY_LT    = "ECEFF1"   # コードブロック背景（ライトグレー）
CODE_HINT  = "E8F5E9"   # ヒントブロック背景（ライトグリーン）
JP_FONT    = "Hiragino Sans"

_bm_id = [2000]

# TOC に載せない見出しテキスト
HEADING_NO_TOC = {"免責事項・著作権表示"}


# ── XML ヘルパー ──────────────────────────────────────────────
def set_style_font(style, *, name=JP_FONT, size=None, color=None, bold=None):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def ensure_style_outline_level(style, level):
    pPr = style.element.get_or_add_pPr()
    for el in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(el)
    ol = OxmlElement("w:outlineLvl")
    ol.set(qn("w:val"), str(level))
    pPr.append(ol)


def enable_mirror_margins(section):
    """見開き印刷用に left_margin/right_margin を「内側/外側」として機能させる。
    これがないと偶数ページでは綴じ側（内側）が外側マージン値になり、
    KDPの内側マージン最低要件（0.375"以上）を満たせなくなる。"""
    sectPr = section._sectPr
    if sectPr.find(qn("w:mirrorMargins")) is None:
        mm = OxmlElement("w:mirrorMargins")
        sectPr.append(mm)


def set_paragraph_outline_level(p, level):
    pPr = p._p.get_or_add_pPr()
    for el in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(el)
    ol = OxmlElement("w:outlineLvl")
    ol.set(qn("w:val"), str(level))
    pPr.append(ol)


def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def para_border(p, *, edges=("left",), color=TEAL, sz="18", space="10"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), space)
        el.set(qn("w:color"), color)
        pBdr.append(el)
    pPr.append(pBdr)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_bookmark(p, name):
    _bm_id[0] += 1
    bid = str(_bm_id[0])
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    p._p.insert(0, start)
    p._p.append(end)


def add_internal_hyperlink(p, anchor, text, *, color=BLUE_DK, bold=False):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), JP_FONT)
    rPr.append(rfonts)
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    if bold:
        rPr.append(OxmlElement("w:b"))
    run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)


def anchor_to_bm(anchor):
    return "bm_" + anchor.strip().lstrip("#").replace("-", "_")


# ── インライン記法 ────────────────────────────────────────────
INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_inline(p, text, *, base_color=INK):
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
            r.font.color.rgb = RGBColor.from_string(base_color)
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
            r.font.color.rgb = RGBColor.from_string(base_color)
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1])
            r.font.name = "Menlo"
            r.font.color.rgb = RGBColor.from_string(LAVENDER_DK)
        elif tok.startswith("[") and "](" in tok:
            label = tok[1:tok.index("](")]
            target = tok[tok.index("](") + 2:-1]
            if target.startswith("#"):
                add_internal_hyperlink(p, anchor_to_bm(target), label)
            else:
                r = p.add_run(label)
                r.font.color.rgb = RGBColor.from_string(BLUE_DK)
        else:
            r = p.add_run(tok)
            r.font.color.rgb = RGBColor.from_string(base_color)


# ── 画像縮小キャッシュ ────────────────────────────────────────
_tmpdir = tempfile.mkdtemp(prefix="mind_book_imgs_")
_img_cache = {}


def prep_image(rel_path):
    if rel_path in _img_cache:
        return _img_cache[rel_path]
    src = os.path.join(BASE_DIR, rel_path)
    if not os.path.isfile(src):
        print(f"  ⚠️  画像が見つかりません: {rel_path}")
        _img_cache[rel_path] = None
        return None
    try:
        im = Image.open(src).convert("RGB")
        long_side = 1100
        w, h = im.size
        scale = min(1.0, long_side / max(w, h))
        if scale < 1.0:
            im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        out = os.path.join(_tmpdir, f"img_{len(_img_cache)}.jpg")
        im.save(out, "JPEG", quality=82, optimize=True)
        result = (out, im.size)
        _img_cache[rel_path] = result
        return result
    except Exception as e:
        print(f"  ⚠️  画像スキップ: {rel_path} ({e})")
        _img_cache[rel_path] = None
        return None


def add_image(doc, rel_path, alt=""):
    prepped = prep_image(rel_path)
    if not prepped:
        return
    path, (w, h) = prepped
    # 5.04"×7.17"トリムサイズで1メソッド=1ページに収めるための小さめサイズ
    max_w, max_h = 1.5, 1.7
    if (w / max_w) >= (h / max_h):
        kwargs = {"width": Inches(max_w)}
    else:
        kwargs = {"height": Inches(max_h)}
    shape = doc.add_picture(path, **kwargs)
    effective_alt = alt if alt else os.path.splitext(os.path.basename(rel_path))[0]
    docPr = shape._inline.find(qn("wp:docPr"))
    if docPr is not None:
        docPr.set("descr", effective_alt)
        docPr.set("title", effective_alt)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(2)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(4)
    if alt:
        cap = doc.add_paragraph(alt)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(6)
        for r in cap.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string("888888")
            r.italic = True


# ── TOC プリスキャン ──────────────────────────────────────────
def build_toc_structure(lines):
    """章見出し（##）と、その配下の各メソッド見出し（### No.XXX）を走査し、
    印刷用の詳細目次（章 → メソッド一覧、ページ番号付き）を組み立てるための
    ネスト構造を返す。メソップのブックマークは本文中の <a id="XXX"> アンカーを
    そのまま再利用するため、Word・Kindle変換のどちらでも同一のリンク先を指す。"""
    structure = []
    current_chapter = None
    chapter_idx = 0
    in_code = False
    in_small = False
    pending_anchor = None
    for raw in lines:
        s = raw.strip()
        if s.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        if s.startswith("<small"):
            in_small = True; continue
        if s.startswith("</small"):
            in_small = False; continue
        if in_small:
            continue
        am = re.match(r'<a id="([^"]+)"></a>', s)
        if am:
            pending_anchor = am.group(1)
            continue
        hm = re.match(r"^(#{1,6})\s+(.*)$", s)
        if not hm:
            continue
        level = len(hm.group(1))
        text = hm.group(2).strip()
        clean = _strip_inline(text)
        if level == 2:
            pending_anchor = None
            if clean == "目次" or clean in HEADING_NO_TOC:
                current_chapter = None
                continue
            chapter_idx += 1
            current_chapter = {"text": clean, "bm": f"chap_{chapter_idx}", "methods": []}
            structure.append(current_chapter)
            continue
        if level == 3:
            is_method = bool(re.match(r'^No\.\d{3}', clean))
            if is_method and current_chapter is not None:
                bm = anchor_to_bm(pending_anchor) if pending_anchor else None
                current_chapter["methods"].append({"text": clean, "bm": bm})
            pending_anchor = None
            continue
        pending_anchor = None
    return structure


def _strip_inline(text):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("**", "").replace("*", "").replace("`", "").strip()


def _toc_pageref_run(tocname):
    r_begin = OxmlElement("w:r")
    rPr_b = OxmlElement("w:rPr")
    ncs = OxmlElement("w:noProof"); rPr_b.append(ncs)
    r_begin.append(rPr_b)
    fc_begin = OxmlElement("w:fldChar"); fc_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(fc_begin)
    r_instr = OxmlElement("w:r")
    rPr_i = OxmlElement("w:rPr")
    ncs2 = OxmlElement("w:noProof"); rPr_i.append(ncs2)
    r_instr.append(rPr_i)
    instr_t = OxmlElement("w:instrText")
    instr_t.set(qn("xml:space"), "preserve")
    instr_t.text = f" PAGEREF {tocname} \\h "
    r_instr.append(instr_t)
    r_sep = OxmlElement("w:r")
    rPr_s = OxmlElement("w:rPr")
    ncs3 = OxmlElement("w:noProof"); rPr_s.append(ncs3)
    r_sep.append(rPr_s)
    fc_sep = OxmlElement("w:fldChar"); fc_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fc_sep)
    r_val = OxmlElement("w:r")
    rPr_v = OxmlElement("w:rPr")
    sz_v = OxmlElement("w:sz"); sz_v.set(qn("w:val"), "20"); rPr_v.append(sz_v)
    r_val.append(rPr_v)
    t_val = OxmlElement("w:t"); t_val.text = "1"; r_val.append(t_val)
    r_end = OxmlElement("w:r")
    rPr_e = OxmlElement("w:rPr")
    ncs4 = OxmlElement("w:noProof"); rPr_e.append(ncs4)
    r_end.append(rPr_e)
    fc_end = OxmlElement("w:fldChar"); fc_end.set(qn("w:fldCharType"), "end")
    r_end.append(fc_end)
    return [r_begin, r_instr, r_sep, r_val, r_end]


def _add_toc_tab(p):
    # ページ幅（5.04"×7.17"トリム）に合わせ、右マージン位置まで正確にドットリーダーを引く
    pos = PAGE_W.twips - MARGIN_INNER.twips - MARGIN_OUTER.twips
    pPr = p._p.get_or_add_pPr()
    tabs_el = pPr.find(qn("w:tabs"))
    if tabs_el is None:
        tabs_el = OxmlElement("w:tabs")
        pPr.append(tabs_el)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(pos))
    tabs_el.append(tab)


def _toc_hyperlink_run(bookmark, text, *, size_half_pt, color, bold=False):
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), bookmark)
    hl.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:eastAsia"), JP_FONT); rPr.append(rf)
    col = OxmlElement("w:color"); col.set(qn("w:val"), color); rPr.append(col)
    if bold:
        rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size_half_pt)); rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t)
    hl.append(run)
    return hl


def emit_detailed_toc(doc, structure):
    """章ごとに全メソッドを一覧化した詳細目次を出力する。
    各行は本文中の見出しブックマークへの実リンク（w:hyperlink）＋ PAGEREF
    フィールドで構成されるため、Word 上でのページ番号更新（Ctrl+A→F9）に
    対応しつつ、Kindle 変換時もクリック可能なナビゲーションとして機能する。"""
    if not structure:
        return
    for chapter in structure:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(10)
        cp.paragraph_format.space_after = Pt(4)
        cp.paragraph_format.keep_with_next = True
        _add_toc_tab(cp)
        cp._p.append(_toc_hyperlink_run(
            chapter["bm"], chapter["text"], size_half_pt=24, color=TEAL_DK, bold=True))
        r_tab = OxmlElement("w:r")
        rPr_tab = OxmlElement("w:rPr")
        sz_tab = OxmlElement("w:sz"); sz_tab.set(qn("w:val"), "24"); rPr_tab.append(sz_tab)
        r_tab.append(rPr_tab)
        r_tab.append(OxmlElement("w:tab"))
        cp._p.append(r_tab)
        for r in _toc_pageref_run(chapter["bm"]):
            cp._p.append(r)

        for method in chapter["methods"]:
            mp = doc.add_paragraph()
            mp.paragraph_format.left_indent = Pt(14)
            mp.paragraph_format.space_after = Pt(2)
            mp.paragraph_format.line_spacing = 1.05
            _add_toc_tab(mp)
            if method["bm"]:
                mp._p.append(_toc_hyperlink_run(
                    method["bm"], method["text"], size_half_pt=18, color=INK))
            else:
                r = mp.add_run(method["text"])
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor.from_string(INK)
            r_tab2 = OxmlElement("w:r")
            rPr_tab2 = OxmlElement("w:rPr")
            sz_tab2 = OxmlElement("w:sz"); sz_tab2.set(qn("w:val"), "18"); rPr_tab2.append(sz_tab2)
            r_tab2.append(rPr_tab2)
            r_tab2.append(OxmlElement("w:tab"))
            mp._p.append(r_tab2)
            if method["bm"]:
                for r in _toc_pageref_run(method["bm"]):
                    mp._p.append(r)


# ── ブロックビルダー ──────────────────────────────────────────
def add_heading(doc, level, text, pending_anchor, page_break=False, toc_bookmark=None):
    stripped_text = _strip_inline(text)
    use_no_toc = (level == 2 and stripped_text in HEADING_NO_TOC)

    if use_no_toc:
        p = doc.add_paragraph(style="Normal")
        if page_break:
            p.paragraph_format.page_break_before = True
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(8)
        add_inline(p, text, base_color=TEAL_DK)
        for r in p.runs:
            r.font.size = Pt(15)
            r.font.bold = True
            rpr = r._element.get_or_add_rPr()
            rf = rpr.get_or_add_rFonts()
            rf.set(qn("w:eastAsia"), JP_FONT)
        para_border(p, edges=("bottom",), color=TEAL, sz="8", space="6")
        if pending_anchor:
            add_bookmark(p, anchor_to_bm(pending_anchor))
        return None

    is_numbered = bool(re.match(r'^\d+-[\dA-Z]', stripped_text))
    if level == 3 and not is_numbered:
        effective_level = 4
    else:
        effective_level = level

    style_map = {1: "Title", 2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}
    outline_map = {2: 0, 3: 1, 4: 2}
    p = doc.add_paragraph(style=style_map[effective_level])
    if page_break:
        p.paragraph_format.page_break_before = True
    if level == 2:
        p.paragraph_format.space_before = Pt(0)
    color_map = {1: TEAL_DK, 2: TEAL_DK, 3: INDIGO_DK, 4: LAVENDER_DK}
    add_inline(p, text, base_color=color_map[level])
    if level == 2:
        para_border(p, edges=("bottom",), color=TEAL, sz="8", space="6")
    if effective_level in outline_map:
        set_paragraph_outline_level(p, outline_map[effective_level])
    if toc_bookmark:
        add_bookmark(p, toc_bookmark)
    if pending_anchor:
        add_bookmark(p, anchor_to_bm(pending_anchor))
    return None


def add_blockquote(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    shade_paragraph(p, TEAL_LT)
    para_border(p, edges=("left",), color=TEAL, sz="24", space="10")
    for i, ln in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        add_inline(p, ln, base_color=INK)


def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    if indent > 0:
        p.paragraph_format.left_indent = Pt(indent * 18)
    add_inline(p, text, base_color=INK)


def add_numbered(doc, num, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    add_inline(p, text, base_color=INK)


def add_paragraph_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    add_inline(p, text, base_color=INK)


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    para_border(p, edges=("bottom",), color=TEAL, sz="6", space="1")


def add_code_block(doc, lines, is_prompt=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Pt(8)
    shade_paragraph(p, CODE_HINT if is_prompt else GRAY_LT)
    border_color = INDIGO if is_prompt else TEAL
    para_border(p, edges=("left",), color=border_color, sz="18", space="10")
    for i, ln in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(ln)
        r.font.name = "Menlo"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(INDIGO_DK if is_prompt else "455A64")
        rpr = r._element.get_or_add_rPr()
        rf = rpr.get_or_add_rFonts()
        rf.set(qn("w:ascii"), "Menlo")
        rf.set(qn("w:hAnsi"), "Menlo")


def add_table(doc, rows):
    header = rows[0]
    body = rows[1:]
    ncol = len(header)
    table = doc.add_table(rows=1 + len(body), cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        shade_cell(cell, TEAL)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(cell_text.strip())
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.size = Pt(10)
    for i, brow in enumerate(body):
        for j in range(ncol):
            cell = table.rows[i + 1].cells[j]
            shade_cell(cell, TEAL_LT if i % 2 == 0 else "FFFFFF")
            cell.paragraphs[0].text = ""
            txt = brow[j].strip() if j < len(brow) else ""
            add_inline(cell.paragraphs[0], txt, base_color=INK)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def split_table_row(line):
    s = line.strip()
    if s.startswith("|"): s = s[1:]
    if s.endswith("|"):   s = s[:-1]
    return [c for c in s.split("|")]


def add_prompt_label(doc, label_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label_text)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(INDIGO_DK)


def add_small_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    add_inline(p, text, base_color="888888")
    for r in p.runs:
        r.font.size = Pt(8)
        rpr = r._element.get_or_add_rPr()
        rf = rpr.get_or_add_rFonts()
        rf.set(qn("w:eastAsia"), JP_FONT)


# ── 用紙サイズ・マージン設定 ──────────────────────────────────
# ── ページ番号フッター ─────────────────────────────────────────
def _page_num_run(char_type=None, instr=None):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz");   sz.set(qn("w:val"), "20");  rPr.append(sz)
    szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "20"); rPr.append(szCs)
    col = OxmlElement("w:color"); col.set(qn("w:val"), "888888"); rPr.append(col)
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:eastAsia"), JP_FONT); rPr.append(rf)
    r.append(rPr)
    if char_type:
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), char_type)
        r.append(fc)
    if instr is not None:
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = instr
        r.append(it)
    return r


def add_page_number_footer(doc):
    for idx, section in enumerate(doc.sections):
        if idx == 0:
            section.different_first_page_header_footer = True

        footer = section.footer
        footer.is_linked_to_previous = False

        if footer.paragraphs:
            p = footer.paragraphs[0]
            for child in list(p._p):
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag in ("r", "hyperlink", "fldSimple"):
                    p._p.remove(child)
        else:
            p = footer.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(0)

        p._p.append(_page_num_run(char_type="begin"))
        p._p.append(_page_num_run(instr=" PAGE "))
        p._p.append(_page_num_run(char_type="separate"))
        r_val = OxmlElement("w:r")
        rPr2 = OxmlElement("w:rPr")
        sz2 = OxmlElement("w:sz"); sz2.set(qn("w:val"), "20"); rPr2.append(sz2)
        r_val.append(rPr2)
        t_val = OxmlElement("w:t"); t_val.text = "1"; r_val.append(t_val)
        p._p.append(r_val)
        p._p.append(_page_num_run(char_type="end"))


# ── 表紙ページ ────────────────────────────────────────────────
def add_cover_page(doc):
    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # メインタイトル（日本語）
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(6)
    tr = tp.add_run("アップルウォッチ・\nマインドフルネス")
    tr.font.name = JP_FONT
    tr.font.size = Pt(19)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor.from_string(TEAL_DK)
    rpr = tr._element.get_or_add_rPr()
    rf  = rpr.get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), JP_FONT)

    # 英語サブタイトル
    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ep.paragraph_format.space_before = Pt(0)
    ep.paragraph_format.space_after  = Pt(12)
    er = ep.add_run("Apple Watch Mindfulness")
    er.font.name = JP_FONT
    er.font.size = Pt(13)
    er.font.bold = False
    er.font.color.rgb = RGBColor.from_string(TEAL)
    rpr_e = er._element.get_or_add_rPr()
    rf_e  = rpr_e.get_or_add_rFonts()
    rf_e.set(qn("w:eastAsia"), JP_FONT)

    # 帯（ティールの罫線）
    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(4)
    div.paragraph_format.space_after  = Pt(12)
    para_border(div, edges=("bottom",), color=TEAL, sz="12", space="4")

    # サブタイトル
    sp2 = doc.add_paragraph()
    sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp2.paragraph_format.space_before = Pt(6)
    sp2.paragraph_format.space_after  = Pt(16)
    sr2 = sp2.add_run("Apple Watchを使ってマインドフルな\n生活を送る100の方法")
    sr2.font.name = JP_FONT
    sr2.font.size = Pt(11)
    sr2.font.bold = True
    sr2.font.color.rgb = RGBColor.from_string(INDIGO_DK)
    rpr2 = sr2._element.get_or_add_rPr()
    rf2  = rpr2.get_or_add_rFonts()
    rf2.set(qn("w:eastAsia"), JP_FONT)

    # アプリ連携の説明
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after  = Pt(20)
    cr = cp.add_run("kmind アプリ × Apple Watch で実現する\n科学的マインドフルネス習慣")
    cr.font.name = JP_FONT
    cr.font.size = Pt(9.5)
    cr.font.italic = True
    cr.font.color.rgb = RGBColor.from_string("546E7A")
    rpr3 = cr._element.get_or_add_rPr()
    rf3  = rpr3.get_or_add_rFonts()
    rf3.set(qn("w:eastAsia"), JP_FONT)

    # 帯（ティールの罫線）
    div2 = doc.add_paragraph()
    div2.paragraph_format.space_before = Pt(4)
    div2.paragraph_format.space_after  = Pt(12)
    para_border(div2, edges=("bottom",), color=TEAL, sz="6", space="4")

    # 著者名
    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ap.paragraph_format.space_before = Pt(0)
    ap.paragraph_format.space_after  = Pt(6)
    ar = ap.add_run("著者：吉田 顕一（Ken Yoshida）")
    ar.font.name = JP_FONT
    ar.font.size = Pt(10.5)
    ar.font.bold = True
    ar.font.color.rgb = RGBColor.from_string(INK)
    rpr4 = ar._element.get_or_add_rPr()
    rf4  = rpr4.get_or_add_rFonts()
    rf4.set(qn("w:eastAsia"), JP_FONT)

    # 発行年
    yp = doc.add_paragraph()
    yp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    yp.paragraph_format.space_before = Pt(0)
    yp.paragraph_format.space_after  = Pt(0)
    yr = yp.add_run("2026")
    yr.font.name = JP_FONT
    yr.font.size = Pt(10)
    yr.font.color.rgb = RGBColor.from_string("888888")
    rpr5 = yr._element.get_or_add_rPr()
    rf5  = rpr5.get_or_add_rFonts()
    rf5.set(qn("w:eastAsia"), JP_FONT)

    # 改ページ
    pb = doc.add_paragraph()
    run_pb = pb.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run_pb._r.append(br)


# ── メイン処理 ────────────────────────────────────────────────
def build():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()

    # ベーススタイル設定（マインドフルネステーマ）
    set_style_font(doc.styles["Normal"],    name=JP_FONT, size=9.5,  color=INK)
    set_style_font(doc.styles["Title"],     name=JP_FONT, size=18,   color=TEAL_DK,   bold=True)
    set_style_font(doc.styles["Heading 1"], name=JP_FONT, size=15,   color=TEAL_DK,   bold=True)
    set_style_font(doc.styles["Heading 2"], name=JP_FONT, size=12,   color=INDIGO_DK, bold=True)
    set_style_font(doc.styles["Heading 3"], name=JP_FONT, size=10.5, color=LAVENDER_DK, bold=True)
    for sname in ("List Bullet", "List Number"):
        try:
            set_style_font(doc.styles[sname], name=JP_FONT, size=9.5, color=INK)
        except KeyError:
            pass
    for hname in ("Heading 1", "Heading 2", "Heading 3"):
        doc.styles[hname].paragraph_format.space_before = Pt(0)
    ensure_style_outline_level(doc.styles["Heading 1"], 0)
    ensure_style_outline_level(doc.styles["Heading 2"], 1)
    ensure_style_outline_level(doc.styles["Heading 3"], 2)

    # 用紙サイズ・マージン設定
    s0 = doc.sections[0]
    s0.page_width         = PAGE_W
    s0.page_height        = PAGE_H
    s0.top_margin         = MARGIN_TOP
    s0.bottom_margin      = MARGIN_BOTTOM
    s0.left_margin        = MARGIN_INNER
    s0.right_margin       = MARGIN_OUTER
    s0.header_distance    = Inches(0.3)
    s0.footer_distance    = Inches(0.3)
    s0.different_first_page_header_footer = True
    enable_mirror_margins(s0)

    # 表紙ページを追加
    add_cover_page(doc)

    toc_structure = build_toc_structure(lines)
    chapter_bm_queue = [c["bm"] for c in toc_structure]

    pending_anchor = None
    in_toc   = False
    in_code  = False
    in_small = False
    code_lines = []
    next_code_is_prompt = False
    skip_first_h1 = True

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # コードフェンス
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines, is_prompt=next_code_is_prompt)
                code_lines = []
                in_code = False
                next_code_is_prompt = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # <small> / </small>
        if stripped.startswith("<small"):
            in_small = True
            i += 1
            continue
        if stripped.startswith("</small"):
            in_small = False
            i += 1
            continue
        if in_small:
            if stripped and stripped != "---":
                add_small_text(doc, stripped)
            i += 1
            continue

        # <div> タグ系はスキップ
        if stripped.startswith("<div") or stripped.startswith("</div"):
            i += 1
            continue

        # アンカー
        m = re.match(r'<a id="([^"]+)"></a>', stripped)
        if m:
            pending_anchor = m.group(1)
            i += 1
            continue

        # 見出し
        hm = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if hm:
            level = min(len(hm.group(1)), 4)
            text = hm.group(2).strip()

            if level == 1 and skip_first_h1:
                skip_first_h1 = False
                i += 1
                continue

            if level == 2 and text == "目次":
                in_toc = True
                add_heading(doc, level, text, pending_anchor, page_break=True)
                emit_detailed_toc(doc, toc_structure)
                pending_anchor = None
                i += 1
                continue
            if level <= 2 and text != "目次":
                in_toc = False
            if in_toc:
                pending_anchor = None
                i += 1
                continue

            is_method_heading = bool(re.match(r'^No\.\d{3}', _strip_inline(text)))
            is_page_break = (level == 2 and text != "免責事項・著作権表示") or \
                            (level == 3 and is_method_heading)
            toc_bm = None
            if level == 2 and text != "目次" and _strip_inline(text) not in HEADING_NO_TOC:
                toc_bm = chapter_bm_queue.pop(0) if chapter_bm_queue else None
            add_heading(doc, level, text, pending_anchor,
                        page_break=is_page_break, toc_bookmark=toc_bm)
            pending_anchor = None
            i += 1
            continue

        # 空行
        if stripped == "":
            i += 1
            continue

        # 目次領域の本文はスキップ
        if in_toc:
            i += 1
            continue

        # 水平線
        if stripped == "---":
            add_divider(doc)
            i += 1
            continue

        # 画像
        im = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if im:
            add_image(doc, im.group(2), im.group(1))
            i += 1
            continue

        # テーブル
        if stripped.startswith("|"):
            tbl = []
            while i < n and lines[i].strip().startswith("|"):
                tbl.append(lines[i])
                i += 1
            parsed = [split_table_row(r) for r in tbl]
            parsed = [r for r in parsed if not all(set(c.strip()) <= set("-: ") and c.strip() for c in r)]
            if parsed:
                add_table(doc, parsed)
            continue

        # 引用
        if stripped.startswith(">"):
            quote = []
            while i < n and lines[i].strip().startswith(">"):
                q = lines[i].strip()[1:].strip()
                quote.append(q.rstrip())
                i += 1
            add_blockquote(doc, quote)
            continue

        # 箇条書き
        bm = re.match(r"^(\s*)[-•]\s+(.+)$", line)
        if bm:
            indent = len(bm.group(1)) // 2
            add_bullet(doc, bm.group(2).strip(), indent=indent)
            i += 1
            continue

        # 番号付きリスト
        nm = re.match(r"^\d+\.\s+(.+)$", stripped)
        if nm:
            add_numbered(doc, 0, nm.group(1).strip())
            i += 1
            continue

        # 通常段落
        add_paragraph_text(doc, stripped)
        i += 1

    # ページ番号フッター
    add_page_number_footer(doc)

    doc.save(OUT)
    size_mb = os.path.getsize(OUT) / 1048576
    print(f"✅ 保存完了: {OUT}  ({size_mb:.1f} MB)")
    print()
    print("【次のステップ】")
    print("  Word で開いて Ctrl+A → F9 でフィールド更新 → 目次のページ番号が確定します")
    print("  KDP へのアップロード前に Word での最終確認を推奨します")


if __name__ == "__main__":
    build()
