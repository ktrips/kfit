#!/usr/bin/env python3
"""
Markdownドキュメントを画像付きWord(.docx)ファイルに変換するスクリプト
- Kindle電子書籍向け A5 フォーマット（デフォルト）
- KDPペーパーバック向け 210x257mm フォーマット（--paperback オプション）
- 各見出しにブックマークを付与
- 目次の各エントリーをそのブックマークへのハイパーリンクとして生成

使い方:
  python3 convert_to_docx.py              # Kindle電子書籍 (A5)
  python3 convert_to_docx.py --paperback  # KDPペーパーバック (210x257mm)
"""

import sys
import os
import re
import io
import unicodedata

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '.pip_pkgs'))

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH  = os.path.join(DOCS_DIR, 'cursor-claude-code-ios-app-book.md')

# ── フォーマット設定 ──────────────────────────────────────────────────────
PAPERBACK_MODE = '--paperback' in sys.argv

if PAPERBACK_MODE:
    # KDP ペーパーバック: 210 x 257 mm、プレミアムカラー用紙（白）
    # KDP推奨マージン: 内側(綴じ側) 19mm、外側 13mm、上下 19mm
    DOCX_PATH    = os.path.join(DOCS_DIR, 'cursor-claude-code-ios-app-book-paperback.docx')
    PAGE_W       = Cm(21.0)
    PAGE_H       = Cm(25.7)
    MARGIN_INNER = Cm(1.9)   # 綴じ側（左）
    MARGIN_OUTER = Cm(1.3)   # 小口側（右）
    MARGIN_TOP   = Cm(1.9)
    MARGIN_BOT   = Cm(1.9)
    BODY_SIZE    = 11.0
    CODE_SIZE    = 9.0
    LINE_SPACING = Pt(20)
    # ── Paperback 画像サイズ（種別ごとに最適化）── 単位: cm（float）──────────
    # テキスト幅 = 210 - 19 - 13 = 178mm
    IMG_WIDTH        = 13.0   # 横長/Web スクリーンショット
    IMG_MAX_H        = 8.0    # 横長の最大高さ
    IMG_W_PHONE      = 9.0    # iPhone 縦長（比率 1.0–1.8）: 300 DPI 相当
    IMG_MAX_H_PHONE  = 13.0   # 縦長の最大高さ
    IMG_W_WATCH      = 5.5    # Apple Watch（小解像度）
    IMG_MAX_H_WATCH  = 7.0
    H_SIZES          = {1: 22, 2: 17, 3: 14, 4: 12}
else:
    # Kindle 電子書籍: A5 (148 x 210 mm)
    DOCX_PATH    = os.path.join(DOCS_DIR, 'cursor-claude-code-ios-app-book.docx')
    PAGE_W       = Cm(14.8)
    PAGE_H       = Cm(21.0)
    MARGIN_INNER = Cm(1.9)
    MARGIN_OUTER = Cm(1.9)
    MARGIN_TOP   = Cm(2.0)
    MARGIN_BOT   = Cm(2.0)
    BODY_SIZE    = 10.5
    CODE_SIZE    = 8.5
    LINE_SPACING = Pt(18)
    IMG_WIDTH        = 9.0    # 横長/Web スクリーンショット（cm float）
    IMG_MAX_H        = 14.0
    IMG_W_PHONE      = 7.0    # iPhone 縦長
    IMG_MAX_H_PHONE  = 12.0
    IMG_W_WATCH      = 4.0    # Apple Watch
    IMG_MAX_H_WATCH  = 5.5
    H_SIZES          = {1: 20, 2: 16, 3: 13, 4: 11}

GREEN_DARK = RGBColor(0x1a, 0x7a, 0x2e)
GREEN_MID  = RGBColor(0x2c, 0x5f, 0x2e)
GRAY       = RGBColor(0x66, 0x66, 0x66)
DARK       = RGBColor(0x22, 0x22, 0x22)
LINK_COLOR = RGBColor(0x1a, 0x7a, 0x2e)

# ─── 画像ユーティリティ ────────────────────────────────────────────────────

def _classify_image(px_w, px_h):
    """
    画像ピクセルサイズから種別と表示サイズを返す。
    戻り値: (target_width_cm, max_height_cm)
    """
    ratio = px_h / px_w if px_w > 0 else 1.0

    if px_w <= 500:
        # Apple Watch スクリーンショット（小型）
        return IMG_W_WATCH, IMG_MAX_H_WATCH
    elif ratio < 1.0:
        # 横長（ウィジェット・Web・Xcodeスクリーンショット）
        return IMG_WIDTH, IMG_MAX_H
    else:
        # 縦長 iPhone スクリーンショット
        return IMG_W_PHONE, IMG_MAX_H_PHONE


def _make_watch_image_bytes(pil_img):
    """
    Apple Watch スクリーンショットに角丸を適用し、
    印刷用に高品質アップスケールして PNG バイト列で返す。
    """
    from PIL import ImageDraw, Image as _PIL

    # 目標: 5.5cm @ 300 DPI ≈ 650 px
    target_px = 650
    w, h = pil_img.size
    scale = target_px / w
    new_w = int(w * scale)
    new_h = int(h * scale)
    pil_img = pil_img.resize((new_w, new_h), _PIL.LANCZOS)

    # 角丸マスク生成
    radius = int(min(new_w, new_h) * 0.14)
    mask = _PIL.new('L', (new_w, new_h), 0)
    ImageDraw.Draw(mask).rounded_rectangle(
        [0, 0, new_w - 1, new_h - 1], radius=radius, fill=255
    )

    # 白背景に合成
    bg = _PIL.new('RGB', (new_w, new_h), (255, 255, 255))
    bg.paste(pil_img.convert('RGB'), mask=mask)

    buf = io.BytesIO()
    bg.save(buf, format='PNG', optimize=True)
    buf.seek(0)
    return buf

# ─── ユーティリティ ────────────────────────────────────────────────────────

def safe_east_asia(run, name='Hiragino Sans'):
    try:
        rPr   = run._element.get_or_add_rPr()
        rFonts= rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), name)
    except Exception:
        pass

def base_run(run, size=None, bold=False, italic=False, color=None, mono=False):
    if size is None:
        size = BODY_SIZE
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if mono:
        run.font.name = 'Courier New'
        safe_east_asia(run, 'Hiragino Sans')
    else:
        run.font.name = 'Hiragino Sans'
        safe_east_asia(run)

def shading(para, fill='F4F4F4'):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)

def setup_styles(doc):
    styles = doc.styles

    def _font(style, size, bold=False, color=None):
        style.font.name   = 'Hiragino Sans'
        style.font.size   = Pt(size)
        style.font.bold   = bold
        if color:
            style.font.color.rgb = color
        try:
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Sans')
        except Exception:
            pass

    _font(styles['Normal'], BODY_SIZE)
    styles['Normal'].paragraph_format.space_after  = Pt(6)
    styles['Normal'].paragraph_format.line_spacing = LINE_SPACING

    for lvl, color in [(1, GREEN_DARK), (2, GREEN_DARK), (3, GREEN_MID), (4, DARK)]:
        h = styles[f'Heading {lvl}']
        _font(h, H_SIZES[lvl], bold=True, color=color)
        h.paragraph_format.space_before    = Pt(6 + lvl * 4)
        h.paragraph_format.space_after     = Pt(8)
        h.paragraph_format.keep_with_next  = True
        # 行間を「倍数」で設定（exact指定だとフォントより小さい場合に文字が切れる）
        # 1.2倍 → フォントサイズに応じて自動的に適切な高さになる
        h.paragraph_format.line_spacing = 1.2

# ─── アンカー生成 ──────────────────────────────────────────────────────────

def make_anchor(text):
    """見出しテキストからユニークなアンカーIDを生成"""
    s = text.strip()
    s = re.sub(r'[「」『』【】「」・…]', '', s)
    s = s.replace('　', '_').replace(' ', '_')
    s = re.sub(r'[^\w\u3040-\u9fff_-]', '', s)
    s = s[:40]
    return s or 'section'

def collect_headings(lines):
    """マークダウン行から (level, text, anchor) のリストを収集"""
    headings = []
    anchor_counts = {}
    in_code = False
    in_toc  = False   # 目次リストはスキップ

    for line in lines:
        s = line.strip()
        if s.startswith('```'):
            in_code = not in_code
            continue
        if in_code:
            continue

        hm = re.match(r'^(#{1,4})\s+(.+)$', s)
        if hm:
            lvl  = len(hm.group(1))
            text = hm.group(2).strip()
            if text == '目次':
                in_toc = True
                continue
            in_toc = False
            base_anchor = make_anchor(text)
            cnt = anchor_counts.get(base_anchor, 0)
            anchor = base_anchor if cnt == 0 else f'{base_anchor}_{cnt}'
            anchor_counts[base_anchor] = cnt + 1
            headings.append((lvl, text, anchor))

    return headings

# ─── ブックマーク ──────────────────────────────────────────────────────────

_bk_counter = [0]

def add_bookmark_to_para(para, anchor):
    """段落にブックマーク開始/終了タグを付与"""
    bk_id = str(_bk_counter[0])
    _bk_counter[0] += 1

    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'),   bk_id)
    start.set(qn('w:name'), anchor)

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bk_id)

    para._p.insert(0, start)
    para._p.append(end)

# ─── 見出し（ブックマーク付き） ───────────────────────────────────────────

def add_heading_with_bookmark(doc, text, level, anchor_map):
    """見出しを追加し、対応するアンカーが存在すればブックマークを付ける"""
    clean = text.strip().lstrip('#').strip()
    lvl   = min(level, 4)
    para  = doc.add_heading(clean, level=lvl)
    anchor = anchor_map.get(clean)
    if anchor:
        add_bookmark_to_para(para, anchor)
    return para

# ─── 目次（ハイパーリンク付き） ───────────────────────────────────────────

def add_hyperlink_para(doc, text, anchor, indent_cm=0.0, size=10.5):
    """アンカーへのハイパーリンクを含む段落を追加"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(indent_cm)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.line_spacing = Pt(16)

    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), anchor)

    r_elem = OxmlElement('w:r')

    rPr = OxmlElement('w:rPr')
    # フォント
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    'Hiragino Sans')
    rFonts.set(qn('w:eastAsia'), 'Hiragino Sans')
    rPr.append(rFonts)
    # サイズ
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    # 色（緑）
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '1a7a2e')
    rPr.append(color_el)
    # 下線
    u_el = OxmlElement('w:u')
    u_el.set(qn('w:val'), 'single')
    rPr.append(u_el)

    r_elem.append(rPr)

    t_elem = OxmlElement('w:t')
    t_elem.text = text
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_elem.append(t_elem)
    hl.append(r_elem)

    para._p.append(hl)
    return para

def _ensure_toc_styles(doc):
    """
    Word標準 TOC 1/2/3 スタイルを確保する。
    これらのスタイルはKDP/Kindleが目次として認識するために必要。
    """
    from docx.enum.style import WD_STYLE_TYPE
    toc_cfg = {
        'TOC 1': (BODY_SIZE + 0.5, 0.0,  False),
        'TOC 2': (BODY_SIZE,       0.6,  False),
        'TOC 3': (BODY_SIZE - 0.5, 1.3,  False),
        'TOC 4': (BODY_SIZE - 1.0, 1.9,  False),
    }
    for name, (size, indent_cm, bold) in toc_cfg.items():
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles['Normal']
        style.font.name = 'Hiragino Sans'
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.left_indent   = Cm(indent_cm)
        style.paragraph_format.space_after   = Pt(3)
        style.paragraph_format.space_before  = Pt(0)
        style.paragraph_format.line_spacing  = Pt(16)
        try:
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Sans')
        except Exception:
            pass

def _add_toc_hyperlink(para, text, anchor, size):
    """TOCエントリーの段落にアンカーハイパーリンクを追加"""
    # テキストからHTMLを除去
    clean_text = re.sub(r'<[^>]+>', '', text).strip()
    if not clean_text:
        return

    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), anchor)

    r_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # フォント
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    'Hiragino Sans')
    rFonts.set(qn('w:eastAsia'), 'Hiragino Sans')
    rPr.append(rFonts)
    # サイズ
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    # 色（緑）
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '1a7a2e')
    rPr.append(color_el)
    # 下線
    u_el = OxmlElement('w:u')
    u_el.set(qn('w:val'), 'single')
    rPr.append(u_el)
    r_elem.append(rPr)

    t_elem = OxmlElement('w:t')
    t_elem.text = clean_text
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_elem.append(t_elem)
    hl.append(r_elem)
    para._p.append(hl)

def generate_toc(doc, headings):
    """
    KDP/Kindle対応の目次を生成。
    Word標準の TOC 1/2/3 段落スタイル＋アンカーハイパーリンクを使用することで
    ・KDPが目次として認識（「目次がありません」警告を解消）
    ・Wordのページレイアウトに従い正しく折り返し
    ・Kindleのチャプターナビゲーションに対応
    """
    SKIP = {'Project Overview', 'Development Commands', 'Safety Rules',
            'Architecture Notes', 'Summary', 'Test plan', 'Notes'}

    filtered = [(lvl, text, anchor) for lvl, text, anchor in headings
                if text not in SKIP]

    # TOC 1/2/3/4 スタイルを確保
    _ensure_toc_styles(doc)

    # ── 目次見出し ──
    toc_heading = doc.add_heading('目次', level=1)
    add_bookmark_to_para(toc_heading, 'toc')

    # ── TOC エントリー（Word標準TOCスタイル＋ハイパーリンク）──
    # TOC 1/2/3 スタイルを使用することでKDPが目次として認識し、
    # ページレイアウトのルールに従って正しく折り返される
    for lvl, text, anchor in filtered:
        if lvl == 1:   toc_style, size = 'TOC 1', BODY_SIZE + 0.5
        elif lvl == 2: toc_style, size = 'TOC 2', BODY_SIZE
        elif lvl == 3: toc_style, size = 'TOC 3', BODY_SIZE - 0.5
        else:          toc_style, size = 'TOC 4', BODY_SIZE - 1.0

        para = doc.add_paragraph(style=toc_style)
        # 折り返し・右マージンを明示設定（はみ出し防止）
        para.paragraph_format.word_wrap    = True
        para.paragraph_format.right_indent = Cm(0.5)
        _add_toc_hyperlink(para, text, anchor, size)

# ─── その他コンテンツ要素 ─────────────────────────────────────────────────

# Word の wp: 名前空間（画像代替テキスト設定用）
_WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

def _set_image_alt_text(run, alt_text):
    """
    Wordの画像に代替テキスト（アクセシビリティ用 alt text）を設定する。
    <wp:docPr descr="..."> 属性として埋め込まれ、スクリーンリーダーや
    KDPのアクセシビリティチェックで利用される。
    """
    if not alt_text:
        return
    for tag in (f'{{{_WP_NS}}}inline', f'{{{_WP_NS}}}anchor'):
        drawing = run._r.find(f'.//{tag}')
        if drawing is not None:
            docPr = drawing.find(f'{{{_WP_NS}}}docPr')
            if docPr is not None:
                docPr.set('descr', alt_text)
                docPr.set('title', alt_text[:100])
            break

def add_image(doc, img_path, caption=None, alt_text=None):
    """
    Word に画像を埋め込む。
    - 非Watchの通常画像: 元ファイルバイトをそのまま埋め込む（再エンコードなし）
    - Apple Watch 画像: 角丸＋高品質アップスケールを適用して PNG で埋め込む
    - 画像種別（iPhone縦長 / Apple Watch / 横長）を自動判定してサイズを決定
    - alt_text: Word 画像の代替テキストとして設定
    """
    if not os.path.exists(img_path):
        p = doc.add_paragraph(f'[画像: {os.path.basename(img_path)}]')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    success = False
    try:
        # Pillow でサイズと種別を取得するだけ（再エンコードはしない）
        pil_img = PILImage.open(img_path)
        px_w, px_h = pil_img.size
        ratio = px_h / px_w if px_w > 0 else 1.0

        target_w_cm, max_h_cm = _classify_image(px_w, px_h)
        target_w = Cm(target_w_cm)
        max_h    = Cm(max_h_cm)

        is_watch = (px_w <= 500)

        if is_watch:
            # Watch: 角丸＋アップスケール処理 → PNG BytesIO
            buf = _make_watch_image_bytes(pil_img)
            # アップスケール後のサイズを再計算
            buf.seek(0)
            upscaled = PILImage.open(buf)
            px_w2, px_h2 = upscaled.size
            ratio = px_h2 / px_w2 if px_w2 > 0 else ratio
            buf.seek(0)
            img_descriptor = buf
        else:
            # 通常画像: 元ファイルバイトを判定して埋め込み
            with open(img_path, 'rb') as f:
                raw = f.read()
            if raw[:2] == b'\xff\xd8':
                # JPEG → そのまま使用（品質劣化なし）
                img_descriptor = io.BytesIO(raw)
            else:
                # PNG → 高品質 JPEG に変換してサイズを削減
                buf = io.BytesIO()
                pil_img.convert('RGB').save(buf, format='JPEG', quality=95, optimize=True)
                buf.seek(0)
                img_descriptor = buf

        # 画像段落（中央揃え・上下スペース）
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.keep_with_next = False
        run = p.add_run()

        # 幅優先でサイズ決定（高さが制約を超えたら高さ基準）
        h_if_width = target_w.cm * ratio
        if h_if_width <= max_h.cm:
            run.add_picture(img_descriptor, width=target_w)
        else:
            run.add_picture(img_descriptor, height=max_h)

        if alt_text:
            _set_image_alt_text(run, alt_text)

        success = True

    except Exception as e:
        import traceback
        traceback.print_exc()
        p = doc.add_paragraph(f'[画像エラー: {os.path.basename(img_path)} — {e}]')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if caption and success:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(12)
        r = cp.runs[0] if cp.runs else cp.add_run(caption)
        base_run(r, size=9, italic=True, color=GRAY)

def add_prompt_label(doc, label_text):
    """
    プロンプト例ラベル段落を追加する。
    下線付き・グレー・小サイズで「▶ プロンプト例: {label_text}」と表示。
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f'▶ プロンプト例: {label_text}')
    run.underline    = True
    run.font.size    = Pt(9.0)
    run.font.color.rgb = GRAY
    safe_east_asia(run)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    shading(p, 'F4F4F4')
    # KDP paperback の「印刷できないマークアップ」警告を防ぐため
    # コード中の < > を全角文字に変換（Swift/TS の型パラメータ対策）
    if PAPERBACK_MODE:
        safe_text = code_text.replace('<', '＜').replace('>', '＞')
    else:
        safe_text = code_text
    r = p.add_run(safe_text)
    base_run(r, size=CODE_SIZE, mono=True)

def add_blockquote(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    shading(p, 'EBF5EC')
    text  = '\n'.join(lines)
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    clean = re.sub(r'`(.+?)`',       r'\1', clean)
    clean = re.sub(r'<[^>]+>',       '',    clean)   # HTMLタグ除去
    r = p.add_run(clean)
    base_run(r, size=10, color=GREEN_MID)

def add_table_from_md(doc, table_lines):
    rows = []
    for line in table_lines:
        if re.match(r'\s*\|[-:| ]+\|\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data[:col_count]):
            cell = row.cells[c_idx]
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            clean = re.sub(r'`(.+?)`',       r'\1', clean)
            clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
            cell.text = clean
            is_header = (r_idx == 0)
            for para in cell.paragraphs:
                for run in para.runs:
                    base_run(run, size=9.5, bold=is_header,
                             color=GREEN_DARK if is_header else None)

def inline_format(para, text):
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'<[^>]+>', '', text)
    pattern = re.compile(r'(\*\*[^*]+?\*\*|`[^`]+`|\*[^*]+?\*)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            base_run(r, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            base_run(r, size=9.5, mono=True)
        elif part.startswith('*') and part.endswith('*'):
            r = para.add_run(part[1:-1])
            base_run(r, italic=True)
        else:
            r = para.add_run(part)
            base_run(r)

def page_break(doc):
    p   = doc.add_paragraph()
    run = p.add_run()
    br  = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

# ─── メイン変換 ───────────────────────────────────────────────────────────

def _add_paperback_titlepage(doc):
    """ペーパーバック用：本文冒頭にタイトルページを挿入"""
    def _centered(text, size, bold=False, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        r = p.add_run(text)
        base_run(r, size=size, bold=bold, color=color)
        return p

    # 上部スペース
    for _ in range(6):
        doc.add_paragraph()

    _centered('AIで週末だけで', 24, bold=True, color=GREEN_DARK, space_after=4)
    _centered('iOS、Apple Watch', 24, bold=True, color=GREEN_DARK, space_after=4)
    _centered('フィットネスアプリを作る方法', 20, bold=True, color=GREEN_DARK, space_after=20)

    # 区切り線代わりの段落
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = sep._p.get_or_add_pPr()
    pBdr= OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    for k, v in [('w:val','single'),('w:sz','6'),('w:space','1'),('w:color','1a7a2e')]:
        bot.set(qn(k), v)
    pBdr.append(bot)
    pPr.append(pBdr)

    doc.add_paragraph()
    _centered('Cursor と Claude で SwiftUI・Apple Health・', 12, color=GRAY, space_after=2)
    _centered('モーションセンサー を動かす個人アプリ開発完全ガイド', 12, color=GRAY, space_after=30)

    for _ in range(4):
        doc.add_paragraph()

    _centered('吉田 顕一（Ken Yoshida）著', 13, bold=True, color=DARK, space_after=4)

    page_break(doc)

def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # ページ設定（Kindle電子書籍 or KDPペーパーバック）
    section = doc.sections[0]
    section.page_width   = PAGE_W
    section.page_height  = PAGE_H
    section.left_margin  = MARGIN_INNER   # 綴じ側
    section.right_margin = MARGIN_OUTER   # 小口側
    section.top_margin   = MARGIN_TOP
    section.bottom_margin= MARGIN_BOT

    setup_styles(doc)

    # ペーパーバック用タイトルページ
    if PAPERBACK_MODE:
        _add_paperback_titlepage(doc)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # ── パス1: 全見出しを収集してアンカーマップを作る ──
    headings   = collect_headings(lines)
    anchor_map = {text: anchor for (_, text, anchor) in headings}
    # TOCに使うリスト（## 目次 自体を除く）
    toc_entries = [(lvl, text, anchor) for (lvl, text, anchor) in headings]

    # ── パス2: 本文を生成 ──
    i = 0
    in_code         = False
    code_type       = ''        # 'text' / 'python' / 'swift' / ...
    in_toc_sec      = False     # ## 目次 ～ 次の## までをスキップ
    code_buf        = []
    table_buf       = []
    pending_prompt  = None      # [プロンプト例]: で読み取ったラベル

    while i < len(lines):
        raw = lines[i].rstrip('\n')
        s   = raw.strip()

        # HTMLブロック要素を一括処理（コードブロック外のみ）
        if not in_code and s.startswith('<') and '>' in s:
            if 'page-break-after' in s:
                # 次の非空行を先読みして見出しレベルを確認
                j = i + 1
                while j < len(lines) and not lines[j].strip():
                    j += 1
                next_s  = lines[j].strip() if j < len(lines) else ''
                next_hm = re.match(r'^(#{1,4})\s+', next_s)
                next_lvl = len(next_hm.group(1)) if next_hm else 0
                # H3/H4（節・項）の前は改ページしない（KindleもPaperbackも）
                # H1/H2（章・部）の前、および本文冒頭のみ改ページする
                if next_lvl not in (3, 4):
                    page_break(doc)
            # それ以外の <div>, <br>, <span> 等はスキップ（Wordに含めない）
            i += 1
            continue

        # [プロンプト例]: description ラベル行（マークダウン独自記法）
        pm = re.match(r'^\[プロンプト例\]:\s*(.+)$', s)
        if pm and not in_code:
            pending_prompt = pm.group(1).strip()
            i += 1
            continue

        # コードブロック
        if s.startswith('```'):
            if not in_code:
                in_code   = True
                code_type = s[3:].strip().lower()   # 'text', 'python', 'swift', ...
                code_buf  = []
                # ```text ブロックの前にプロンプトラベルを追加
                if code_type == 'text':
                    lbl = pending_prompt if pending_prompt else '（実装・調査プロンプトの例）'
                    add_prompt_label(doc, lbl)
                pending_prompt = None
            else:
                in_code   = False
                code_type = ''
                add_code_block(doc, '\n'.join(code_buf))
            i += 1
            continue
        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # 見出し
        hm = re.match(r'^(#{1,4})\s+(.+)$', s)
        if hm:
            lvl  = len(hm.group(1))
            text = hm.group(2).strip()

            # ## 目次 セクション開始 → 目次を生成してスキップ開始
            if text == '目次':
                in_toc_sec = True
                generate_toc(doc, toc_entries)
                i += 1
                continue

            # 目次内のリスト行は in_toc_sec フラグで後処理
            if in_toc_sec:
                in_toc_sec = False  # 次の見出しで目次セクション終了

            add_heading_with_bookmark(doc, text, lvl, anchor_map)
            i += 1
            continue

        # 目次セクション内のリスト行（- [xxx](#yyy)）はスキップ
        if in_toc_sec:
            i += 1
            continue

        # テーブル行の収集
        if s.startswith('|'):
            table_buf.append(s)
            i += 1
            next_s = lines[i].strip() if i < len(lines) else ''
            if not next_s.startswith('|'):
                add_table_from_md(doc, table_buf)
                table_buf = []
                doc.add_paragraph()
            continue
        if table_buf:
            add_table_from_md(doc, table_buf)
            table_buf = []
            doc.add_paragraph()

        # 水平線
        if re.match(r'^---+\s*$', s):
            p   = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr= OxmlElement('w:pBdr')
            bot = OxmlElement('w:bottom')
            for k, v in [('w:val','single'),('w:sz','4'),('w:space','1'),('w:color','AAAAAA')]:
                bot.set(qn(k), v)
            pBdr.append(bot)
            pPr.append(pBdr)
            i += 1
            continue

        # 画像
        im = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', s)
        if im:
            alt_text = im.group(1).strip()   # マークダウンの ![代替テキスト]
            rel      = im.group(2)
            abs_     = os.path.join(DOCS_DIR, rel)
            caption  = None
            if i + 1 < len(lines):
                nxt = lines[i+1].strip()
                cm  = re.match(r'^\*▲\s*(.+)\*$', nxt)
                if cm:
                    caption = '▲ ' + cm.group(1)
                    i += 1
            # Paperback: 明示キャプション（*▲*）がない場合、
            # マークダウンの alt text をキャプションとして表示する
            if PAPERBACK_MODE and not caption and alt_text:
                caption = '▲ ' + alt_text
            add_image(doc, abs_, caption, alt_text=alt_text)
            i += 1
            continue

        # 引用ブロック
        if s.startswith('>'):
            q_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                ql = lines[i].strip().lstrip('>').strip()
                if ql:
                    q_lines.append(ql)
                i += 1
            if q_lines:
                add_blockquote(doc, q_lines)
            continue

        # リスト
        lm = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', s)
        if lm:
            indent  = len(lm.group(1))
            marker  = lm.group(2)
            content = lm.group(3)
            style   = 'List Bullet' if not marker[0].isdigit() else 'List Number'
            p       = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.3)
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
            clean = re.sub(r'`(.+?)`',       r'\1', clean)
            clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
            clean = re.sub(r'<[^>]+>',               '',    clean)  # HTMLタグ除去
            r = p.add_run(clean)
            base_run(r)
            i += 1
            continue

        # キャプション行（画像処理済みのもの）
        if re.match(r'^\*▲.+\*$', s):
            i += 1
            continue

        # 空行
        if not s:
            i += 1
            continue

        # 通常段落
        p = doc.add_paragraph()
        inline_format(p, s)
        i += 1

    doc.save(docx_path)
    mode_label = 'KDPペーパーバック (210×257mm)' if PAPERBACK_MODE else 'Kindle電子書籍 (A5)'
    print(f'✅ Word文書を保存しました [{mode_label}]')
    print(f'   パス: {docx_path}')
    print(f'   ファイルサイズ: {os.path.getsize(docx_path) // 1024} KB')
    print(f'   収録見出し数: {len(headings)} 件')
    if PAPERBACK_MODE:
        print(f'   画像最大幅: {IMG_WIDTH}cm / 最大高さ: {IMG_MAX_H}cm')

if __name__ == '__main__':
    convert_md_to_docx(MD_PATH, DOCX_PATH)
