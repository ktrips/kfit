#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
cursor-claude-code-ios-app-book-plus.docx → Kindle Paperback (210×257mm) 変換スクリプト

処理内容:
  1) 用紙サイズを 210 × 257 mm (8.27 × 10.11 インチ) に変更
  2) KDP推奨マージンを設定（内側 19mm / 外側 13mm / 上下 13mm）
  3) フッターにページ番号を追加（中央揃え）
  4) 目次フィールドはそのまま保持（Wordで開いて「フィールド更新」で反映）

pip install python-docx
"""

import copy
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = os.path.join(os.path.dirname(__file__), "cursor-claude-code-ios-app-book-plus.docx")
OUT = SRC  # 上書き保存

# ── Kindle Paperback 8.27 × 10.11 インチ = 210 × 257 mm ──
PAGE_W = Mm(210)
PAGE_H = Mm(257)

# KDP推奨マージン（ページ数が多い場合は内側を広めに）
MARGIN_TOP    = Mm(19)
MARGIN_BOTTOM = Mm(19)
MARGIN_INNER  = Mm(19)   # 内側（ノド）
MARGIN_OUTER  = Mm(13)   # 外側（小口）


def twips(emu):
    """EMU → twips 変換（1 twip = 635 EMU）"""
    return str(int(emu / 635))


def set_page_size_and_margins(doc):
    """全セクションの用紙サイズ・マージンを変更する"""
    for section in doc.sections:
        # 用紙サイズ
        section.page_width  = PAGE_W
        section.page_height = PAGE_H

        # マージン
        section.top_margin    = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin   = MARGIN_INNER   # 右綴じ(日本語)は左=内側
        section.right_margin  = MARGIN_OUTER

        # ヘッダー/フッターの距離
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)


def add_page_number_footer(doc):
    """全セクションのフッター中央にページ番号フィールドを追加する"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # 既存段落をクリアまたは再利用
        if footer.paragraphs:
            p = footer.paragraphs[0]
            # 段落内のrunをすべて削除
            for child in list(p._p):
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag in ("r", "hyperlink", "fldSimple"):
                    p._p.remove(child)
        else:
            p = footer.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(0)

        # --- ページ番号フィールド: { PAGE } ---
        def fld_run(char_type=None, instr=None):
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "20")   # 10pt
            rPr.append(sz)
            szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "20")
            rPr.append(szCs)
            r.append(rPr)
            if char_type:
                fc = OxmlElement("w:fldChar")
                fc.set(qn("w:fldCharType"), char_type)
                r.append(fc)
            if instr is not None:
                it = OxmlElement("w:instrText")
                it.set(qn("xml:space"), "preserve")
                it.text = instr
                r.append(it)
            return r

        p._p.append(fld_run(char_type="begin"))
        p._p.append(fld_run(instr=" PAGE "))
        p._p.append(fld_run(char_type="separate"))
        # プレースホルダー（Wordが自動更新する）
        r_val = OxmlElement("w:r")
        rPr2 = OxmlElement("w:rPr")
        sz2 = OxmlElement("w:sz"); sz2.set(qn("w:val"), "20"); rPr2.append(sz2)
        r_val.append(rPr2)
        t_val = OxmlElement("w:t"); t_val.text = "1"; r_val.append(t_val)
        p._p.append(r_val)
        p._p.append(fld_run(char_type="end"))


def ensure_toc_page_numbering(doc):
    """
    目次フィールド (TOC) にページ番号表示スイッチが含まれているか確認し、
    なければ \\p スイッチを追加する。
    ※ 既に \\h \\z \\u が付いている場合はそのまま（Wordで更新時に機能する）
    """
    for p in doc.paragraphs:
        for instr in p._p.iter(qn("w:instrText")):
            if instr.text and "TOC" in instr.text:
                # \\n スイッチ（ページ番号非表示）が付いている場合は除去
                if r"\n" in instr.text:
                    instr.text = instr.text.replace(r"\n", "").strip()
                # \\p（リーダー）が無ければ追加しない（KDPはリーダー不要）
                break


def main():
    print(f"読み込み中: {SRC}")
    doc = Document(SRC)

    print("用紙サイズ・マージンを設定中 (210 × 257 mm)...")
    set_page_size_and_margins(doc)

    print("フッターにページ番号を追加中...")
    add_page_number_footer(doc)

    print("目次フィールドを確認中...")
    ensure_toc_page_numbering(doc)

    doc.save(OUT)
    size_mb = os.path.getsize(OUT) / 1048576
    print(f"✅ 保存完了: {OUT}  ({size_mb:.1f} MB)")
    print()
    print("【次のステップ】")
    print("  1) Word で出力ファイルを開く")
    print("  2) Ctrl+A (全選択) → F9 (フィールド更新) → 目次のページ番号が反映されます")
    print("  3) KDP へアップロード前に PDF エクスポートを推奨")


if __name__ == "__main__":
    main()
