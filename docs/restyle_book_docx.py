#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
既存の AppleWatchDiet_100methods_PB.docx を「フォーマットを変えずに」加工する。
  1) 目次をKindleの目次として使えるように：
     - 見出し段落にアウトラインレベル(w:outlineLvl)を付与 → KindleがナビゲーションTOC(NCX)を自動生成
     - 各見出しにブックマークを設定
     - 既存の目次の各行をブックマークへのハイパーリンク化（文字装飾は変更しない＝見た目そのまま）
  2) すべての画像に代替テキスト(descr)を付与（既存のものは保持）
  3) 配色を青基調 → 黄緑基調へ（色の値だけ置換。レイアウト・装飾は不変）。オレンジはそのまま維持。
入力ファイルを上書き保存する（事前にバックアップ済み）。
"""
import re
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PATH = "AppleWatchDiet_100methods_PB.docx"

# ── 配色マップ（青系 → 黄緑系。オレンジ #F57C00 / クリーム #FFF8E1 は不変） ──
# テキスト色（w:color/@w:val）：濃いめの黄緑で可読性を確保
TEXT_MAP = {
    "1A73E8": "558B2F",  # 主役の青文字 → 濃い黄緑
    "0D47A1": "33691E",  # 濃紺文字 → 深緑
}
# 塗り（w:shd/@w:fill, 罫線/@w:color など）：白文字が乗る面は少し濃い黄緑に
FILL_MAP = {
    "1A73E8": "689F38",  # 青カード背景 → 黄緑（白文字が読める濃さ）
    "0D47A1": "33691E",  # 濃紺バンド → 深緑
    "E3F2FD": "F1F8E9",  # 薄い水色背景 → 薄い黄緑
}

GENERIC_ALT = "Apple Watchダイエットの解説イラスト"

# ── 見出しトップセクション（正規化テキスト → (level, bookmark名)） ──
TOP_SECTIONS = {
    "はじめに": (0, "s_intro"),
    "AppleWatchとFitingoとは": (0, "s_about"),
    "目次": (0, "s_toc"),
    "100のダイエット方法": (0, "s_methods"),
    "Fitingoアプリ20の機能": (0, "s_features"),
    "📅平日のルーティン": (0, "s_weekday"),
    "🏆週末のルーティン（強化モード）": (0, "s_weekend"),
    "おわりに": (0, "s_closing"),
    "著者プロフィール": (0, "s_author"),
}


def norm(t):
    return re.sub(r"[\s\u3000]", "", t)


# ── ヘルパー ───────────────────────────────────────────────
def set_outline_level(p, level):
    pPr = p._p.get_or_add_pPr()
    for e in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(e)
    ol = OxmlElement("w:outlineLvl")
    ol.set(qn("w:val"), str(level))
    rpr = pPr.find(qn("w:rPr"))
    sect = pPr.find(qn("w:sectPr"))
    ref = rpr if rpr is not None else sect
    if ref is not None:
        ref.addprevious(ol)
    else:
        pPr.append(ol)


def add_bookmark(p, name, bid):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(start)
    else:
        p._p.insert(0, start)
    p._p.append(end)


def wrap_runs_in_hyperlink(p, anchor):
    """段落内の w:r を w:hyperlink で囲む（rPr等の装飾は維持＝見た目不変）。"""
    pEl = p._p
    runs = pEl.findall(qn("w:r"))
    if not runs:
        return False
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), anchor)
    hl.set(qn("w:history"), "1")
    runs[0].addprevious(hl)
    for r in runs:
        hl.append(r)
    return True


def recolor(root):
    n = 0
    for el in root.iter():
        tag = el.tag
        if tag == qn("w:color"):
            v = el.get(qn("w:val"))
            if v and v.upper() in TEXT_MAP:
                el.set(qn("w:val"), TEXT_MAP[v.upper()]); n += 1
        elif tag == qn("w:shd"):
            f = el.get(qn("w:fill"))
            if f and f.upper() in FILL_MAP:
                el.set(qn("w:fill"), FILL_MAP[f.upper()]); n += 1
            c = el.get(qn("w:color"))
            if c and c.upper() in FILL_MAP:
                el.set(qn("w:color"), FILL_MAP[c.upper()]); n += 1
        else:
            c = el.get(qn("w:color"))
            if c and c.upper() in FILL_MAP:
                el.set(qn("w:color"), FILL_MAP[c.upper()]); n += 1
            f = el.get(qn("w:fill"))
            if f and f.upper() in FILL_MAP:
                el.set(qn("w:fill"), FILL_MAP[f.upper()]); n += 1
    return n


def main():
    doc = Document(PATH)
    ps = doc.paragraphs

    # 既存 bookmark id の最大値を調べる
    max_id = 0
    for bm in doc.element.iter(qn("w:bookmarkStart")):
        try:
            max_id = max(max_id, int(bm.get(qn("w:id"))))
        except (TypeError, ValueError):
            pass
    bid = max(max_id + 1, 2000)

    # ── 1) 見出し：アウトラインレベル + ブックマーク ──
    marker_re = re.compile(r"^No\.\s+(\d{3})\s*$")
    headings = 0
    methods = 0
    i = 0
    while i < len(ps):
        p = ps[i]
        t = p.text.strip()
        key = norm(t)
        if key in TOP_SECTIONS:
            level, name = TOP_SECTIONS[key]
            set_outline_level(p, level)
            add_bookmark(p, name, bid); bid += 1
            headings += 1
            i += 1
            continue
        m = marker_re.match(t)
        if m:
            num = m.group(1)
            # 次の非空段落（タイトル行）を見出しにする
            j = i + 1
            while j < len(ps) and not ps[j].text.strip():
                j += 1
            target = ps[j] if j < len(ps) else p
            set_outline_level(target, 1)
            add_bookmark(target, f"m{num}", bid); bid += 1
            methods += 1
            i = j + 1
            continue
        i += 1

    # ── 2) 既存の目次行をハイパーリンク化（見た目は変えない） ──
    toc_no_re = re.compile(r"^No\.(\d{3})\b")          # "No.001 …… 15"（目次専用：スペース無し）
    links = 0
    for idx, p in enumerate(ps):
        t = p.text.strip()
        if not t:
            continue
        anchor = None
        # 章レベルの目次（先頭付近、ページ番号付き）
        if idx < 33 and re.search(r"[…\.]{2,}\s*\d+\s*$", t):
            nk = norm(t)
            if "はじめに" in nk:
                anchor = "s_intro"
            elif "AppleWatchとFitingo" in nk:
                anchor = "s_about"
            elif "20の機能" in nk:
                anchor = "s_features"
            elif "平日" in nk:
                anchor = "s_weekday"
            elif "週末" in nk:
                anchor = "s_weekend"
        # 方法インデックス "No.001 …… 15"
        elif toc_no_re.match(t):
            anchor = "m" + toc_no_re.match(t).group(1)
        # おわりに / 著者プロフィール のインデックス行
        elif re.search(r"[…\.]{2,}\s*\d+\s*$", t):
            if t.startswith("おわりに"):
                anchor = "s_closing"
            elif t.startswith("著者プロフィール"):
                anchor = "s_author"
        if anchor and wrap_runs_in_hyperlink(p, anchor):
            links += 1

    # ── 3) 画像の代替テキスト（既存は保持、文脈から補完） ──
    alt_set = 0
    last_text = ""
    for p in doc.paragraphs:
        drawings = p._p.findall(f".//{qn('w:drawing')}")
        t = p.text.strip()
        if drawings:
            alt = (t or last_text or GENERIC_ALT)[:120]
            for dr in drawings:
                for docPr in dr.iter(qn("wp:docPr")):
                    if not docPr.get("descr"):
                        docPr.set("descr", alt); alt_set += 1
                    if not docPr.get("title"):
                        docPr.set("title", alt[:80])
        if t:
            last_text = t
    # 取りこぼし（表内・段落外など）を汎用テキストで補完
    for docPr in doc.element.iter(qn("wp:docPr")):
        if not docPr.get("descr"):
            docPr.set("descr", GENERIC_ALT); alt_set += 1
        if not docPr.get("title"):
            docPr.set("title", GENERIC_ALT)

    # ── 4) 配色置換（本文＋スタイル） ──
    c1 = recolor(doc.element)
    c2 = recolor(doc.styles.element)

    doc.save(PATH)
    print(f"見出し(トップ):{headings}  方法:{methods}  目次リンク:{links}")
    print(f"代替テキスト付与:{alt_set}")
    print(f"配色置換: 本文{c1} / スタイル{c2}")


if __name__ == "__main__":
    main()
